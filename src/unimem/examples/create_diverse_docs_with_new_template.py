***REMOVED***!/usr/bin/env python3
"""
基于新模板创建5份差异化的短视频需求文档

使用新的智能填充模板，创建5个不同场景的需求文档。
"""

import sys
import os
from pathlib import Path
from datetime import datetime

***REMOVED*** 添加项目路径
project_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(project_root))

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_diverse_doc(output_path: str, scenario: dict) -> str:
    """
    基于新模板格式创建差异化的需求文档
    
    Args:
        output_path: 输出文件路径
        scenario: 场景配置字典
        
    Returns:
        创建的文档路径
    """
    doc = Document()
    
    ***REMOVED*** 标题
    title = doc.add_heading('短视频创作需求模板', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    ***REMOVED*** 说明（新模板格式）
    doc.add_paragraph('请根据以下说明填写相应内容，所有内容都是可选的，根据需要填写即可。')
    doc.add_paragraph('💡 提示：建议填写尽可能详细的信息，以便生成更符合您需求的视频剧本。')
    doc.add_paragraph()
    
    ***REMOVED*** 1. 视频基本信息
    doc.add_heading('一、视频基本信息', level=1)
    doc.add_paragraph('视频类型（必填，从以下选项选择其一）：')
    doc.add_paragraph('  - ecommerce（电商推广）', style='List Bullet')
    doc.add_paragraph('  - ip_building（个人IP打造）', style='List Bullet')
    doc.add_paragraph('  - knowledge（知识分享）', style='List Bullet')
    doc.add_paragraph('  - vlog（生活Vlog）', style='List Bullet')
    doc.add_paragraph('  - media（自媒体内容）', style='List Bullet')
    doc.add_paragraph(f'视频类型: {scenario["video_type"]}')
    doc.add_paragraph()
    
    doc.add_paragraph('目标平台（必填，从以下选项选择其一）：')
    doc.add_paragraph('  - douyin（抖音）', style='List Bullet')
    doc.add_paragraph('  - xiaohongshu（小红书）', style='List Bullet')
    doc.add_paragraph('  - tiktok（TikTok国际）', style='List Bullet')
    doc.add_paragraph('  - youtube（YouTube）', style='List Bullet')
    doc.add_paragraph(f'目标平台: {scenario["platform"]}')
    doc.add_paragraph()
    
    doc.add_paragraph(f'目标时长（秒，可选，默认60秒）: {scenario.get("duration", 60)}')
    doc.add_paragraph()
    
    ***REMOVED*** 2. 当前任务需求
    doc.add_heading('二、当前任务需求', level=1)
    doc.add_paragraph('请详细描述本次视频创作的具体需求，每行一条：')
    for req in scenario.get("task_memories", []):
        doc.add_paragraph(req)
    doc.add_paragraph()
    
    ***REMOVED*** 3. 修改需求（可选）
    doc.add_heading('三、修改需求（可选）', level=1)
    doc.add_paragraph('如果在已有脚本基础上进行修改，请填写修改要求：')
    for mod in scenario.get("modification_memories", []):
        doc.add_paragraph(mod)
    if not scenario.get("modification_memories"):
        doc.add_paragraph('（首次生成，无需填写）')
    doc.add_paragraph()
    
    ***REMOVED*** 4. 通用记忆总结（可选）
    doc.add_heading('四、通用记忆总结（可选）', level=1)
    doc.add_paragraph('跨任务的用户偏好和通用风格偏好，这些会应用到所有视频创作：')
    for mem in scenario.get("general_memories", []):
        doc.add_paragraph(mem)
    doc.add_paragraph()
    
    ***REMOVED*** 5. 用户偏好设置
    doc.add_heading('五、用户偏好设置（可选）', level=1)
    doc.add_paragraph('请使用"键: 值"的格式填写，例如：')
    for key, value in scenario.get("user_preferences", {}).items():
        doc.add_paragraph(f'{key}: {value}')
    doc.add_paragraph()
    
    ***REMOVED*** 6. 商品信息（仅电商题材需要）
    if scenario.get("product_info"):
        doc.add_heading('六、商品信息（仅电商题材需要）', level=1)
        doc.add_paragraph('请使用"键: 值"的格式填写，例如：')
        for key, value in scenario["product_info"].items():
            doc.add_paragraph(f'{key}: {value}')
        doc.add_paragraph()
    
    ***REMOVED*** 7. 镜头素材（新格式，动态生成）
    doc.add_heading('七、镜头素材', level=1)
    doc.add_paragraph('请描述可用的镜头素材，每行一个镜头，使用"键: 值"格式或直接描述：')
    doc.add_paragraph('💡 提示：镜头素材越详细，生成的剧本越精准。建议包括：产品展示、使用场景、细节特写、对比效果等。')
    
    shot_materials = scenario.get("shot_materials", [])
    if shot_materials:
        for shot in shot_materials:
            doc.add_paragraph(shot)
    else:
        ***REMOVED*** 根据视频类型生成默认镜头建议
        if scenario["video_type"] == "ecommerce":
            default_shots = [
                "镜头1: 产品整体展示（全景，展示产品全貌）",
                "镜头2: 产品特写（细节展示，突出卖点）",
                "镜头3: 使用场景展示（实际应用场景）",
                "镜头4: 对比效果（使用前后对比）",
                "镜头5: 用户反应镜头（惊喜/满意表情）",
                "镜头6: 产品细节特写（材质/做工细节）"
            ]
        elif scenario["video_type"] == "knowledge":
            default_shots = [
                "镜头1: 标题和讲师介绍（开场，建立权威性）",
                "镜头2: 核心概念展示（动画/图示，可视化解释）",
                "镜头3: 实例演示（实际操作，真实案例）",
                "镜头4: 重点知识点标注（强调关键信息）",
                "镜头5: 总结和回顾（结尾，关键点回顾）"
            ]
        else:
            default_shots = [
                "镜头1: 开场镜头（吸引注意，建立情境）",
                "镜头2: 主要内容展示（核心内容呈现）",
                "镜头3: 细节特写（重点内容强调）",
                "镜头4: 结尾总结（收尾，引导行动）"
            ]
        
        for shot in default_shots[:scenario.get("duration", 60) // 10]:
            doc.add_paragraph(shot)
    
    doc.add_paragraph()
    
    ***REMOVED*** 保存文档
    doc.save(output_path)
    return output_path


def main():
    """主函数：创建5份差异化的需求文档"""
    
    ***REMOVED*** 定义5个差异化场景（每个至少30秒，剧本至少300字）
    scenarios = [
        {
            "id": "scenario_1",
            "name": "电商-美妆-小红书-简单",
            "video_type": "ecommerce",
            "platform": "xiaohongshu",
            "duration": 45,  ***REMOVED*** 至少30秒
            "task_memories": [
                "推广新品春季限定口红，这是一款专为年轻女性设计的彩妆产品，适合日常通勤、约会聚会等多种场合使用",
                "核心卖点：12小时超长持久不脱色，色彩饱满显色度高，柔雾质地高级感十足，质地舒适不拔干，包装精致适合送礼",
                "目标受众：18-28岁年轻女性，喜欢清新自然妆容风格，追求性价比和实用性，注重产品实际效果",
                "风格要求：清新自然，真实体验分享，避免过度美颜滤镜，保持真实感，让观众看到产品的真实状态",
                "需要详细展示：试色效果（自然光线、室内光线、室外光线的显色度对比）、持久度测试（时间轴展示0小时、6小时、12小时后的效果对比）、多色号对比（5个色号的详细对比和不同肤色适配建议）、日常妆容搭配技巧（通勤、约会、聚会等不同场合的妆容效果）",
                "强调产品优势：持久不脱色（12小时超长持久，吃饭喝水不掉色）、显色度高（色彩饱满，一次涂抹即可显色）、质地舒适（柔雾质地，不拔干不起皮）、包装精致（适合送礼，性价比高）",
                "内容要求：真实使用体验分享，生活化语言表达，避免过度营销感，突出产品实际效果，让观众产生购买欲望",
                "互动要求：引导观众分享自己的口红使用心得，增加互动性和转化率，鼓励评论和提问"
            ],
            "modification_memories": [],
            "general_memories": [
                "喜欢用生活化语言，像朋友分享一样自然",
                "偏好真实体验分享，展示实际使用效果",
                "避免过度美颜滤镜，保持真实感",
                "注重细节展示，让观众看到产品的真实状态"
            ],
            "user_preferences": {
                "风格偏好": "清新自然，真实分享",
                "语气偏好": "亲切友好，像朋友推荐",
                "平台偏好": "小红书",
                "内容偏好": "详细实用，有参考价值"
            },
            "product_info": {
                "产品名称": "春季限定柔雾口红",
                "核心卖点": "12小时超长持久，色彩饱满显色，柔雾质地高级感",
                "价格区间": "100-150元",
                "色号": "5个色号可选（豆沙色、玫瑰色、珊瑚色、正红色、梅子色）",
                "适用场景": "日常通勤、约会、聚会、拍照"
            },
            "shot_materials": [
                "镜头1: 产品整体展示（口红外观，精致包装设计，品牌logo特写）",
                "镜头2: 产品特写（膏体细节，色彩展示，质地质感）",
                "镜头3: 试色过程（自然光线下的试色效果，室内外光线对比）",
                "镜头4: 持久度测试（时间轴展示：0小时、6小时、12小时后的效果对比）",
                "镜头5: 多色号展示（5个色号的对比，不同肤色适配建议）",
                "镜头6: 日常妆容搭配（不同场合的妆容效果：通勤、约会、聚会）",
                "镜头7: 使用技巧分享（如何涂抹更持久，如何搭配不同妆容风格）",
                "镜头8: 包装细节展示（精致包装，适合送礼，性价比说明）"
            ]
        },
        {
            "id": "scenario_2",
            "name": "教育-编程-抖音-中等",
            "video_type": "knowledge",
            "platform": "douyin",
            "duration": 60,  ***REMOVED*** 至少30秒
            "task_memories": [
                "Python编程入门教学，面向零基础初学者，帮助用户从零开始学习Python编程语言",
                "核心内容：详细讲解变量和数据类型，包括整数、浮点数、字符串、布尔值等基础概念，以及变量命名规则、数据类型转换、变量赋值等实际操作",
                "目标受众：编程初学者，对Python感兴趣但没有任何编程基础的用户，希望通过视频学习掌握Python基础",
                "风格要求：通俗易懂，循序渐进，用生活化的例子解释编程概念，让非专业用户也能理解",
                "教学重点：什么是变量（用生活例子解释）、为什么需要变量（存储数据的重要性）、如何命名变量（命名规则和最佳实践）、不同数据类型的区别和使用场景（整数、浮点数、字符串、布尔值的详细对比）",
                "需要详细讲解：变量定义和赋值（实际代码演示）、数据类型转换（如何在不同类型间转换）、实际应用场景（计算器、数据存储、用户输入等真实案例）",
                "内容要求：代码演示清晰易懂，增加详细注释说明每一行代码的作用，实例贴近实际应用场景，讲解节奏适中不拖沓，让初学者容易跟上和理解",
                "互动要求：设置思考题和练习题，引导观众动手实践编写代码，鼓励提问和反馈，建立学习社区氛围"
            ],
            "modification_memories": [],
            "general_memories": [
                "喜欢循序渐进的教学方式，从简单到复杂",
                "偏好实际案例演示，用真实项目场景讲解",
                "避免过于理论化，注重实践操作",
                "强调代码注释的重要性，让代码更易理解"
            ],
            "user_preferences": {
                "风格偏好": "专业但易懂，深入浅出",
                "语气偏好": "耐心细致，像老师一样",
                "平台偏好": "抖音",
                "内容偏好": "实用性强，可以立即应用"
            },
            "product_info": {},
            "shot_materials": [
                "镜头1: 课程标题和讲师介绍（Python编程入门，建立权威性）",
                "镜头2: 开场引入（用生活例子解释什么是变量，降低理解门槛）",
                "镜头3: 代码演示屏幕录制（变量定义和赋值，详细注释说明）",
                "镜头4: 实际运行效果展示（代码执行结果，输出展示）",
                "镜头5: 重点知识点标注（变量命名规则，关键概念强调）",
                "镜头6: 实例演示（贴近实际应用场景的案例：计算器、数据存储等）",
                "镜头7: 注释说明展示（代码注释详解，为什么这样写）",
                "镜头8: 练习题展示（巩固练习，互动环节，引导动手实践）",
                "镜头9: 常见错误提醒（初学者容易犯的错误，如何避免）",
                "镜头10: 总结和回顾（本节重点回顾，下节预告）"
            ]
        },
        {
            "id": "scenario_3",
            "name": "娱乐-搞笑-抖音-简单",
            "video_type": "media",
            "platform": "douyin",
            "duration": 45,  ***REMOVED*** 至少30秒
            "task_memories": [
                "创作搞笑日常段子，展示生活中的小趣事，让观众在轻松愉快的氛围中产生共鸣和笑声",
                "目标受众：18-35岁年轻人，喜欢轻松幽默的内容，追求真实有趣的生活分享，喜欢互动和评论",
                "风格要求：轻松幽默，真实自然，避免过度夸张，保持生活化场景，让观众有代入感",
                "内容要求：开头3秒内制造悬念吸引注意，中间部分笑点突出有节奏感，结尾有记忆点让观众印象深刻，整体节奏紧凑不拖沓",
                "剧情要求：日常场景建立情境（让观众有代入感）、意外情况制造转折（剧情推进的关键点）、反应镜头表达情绪（表情动作要自然）、反转或笑点制造惊喜（高潮部分要出彩）",
                "表现要求：自然表演不做作，真实生活场景拍摄，避免过度夸张失真，保持真实感和生活气息，让观众产生共鸣",
                "互动要求：引导观众评论分享自己的类似经历，增加互动性和参与感，鼓励点赞转发，建立粉丝粘性",
                "时长控制：45秒内完成完整剧情，前3秒抓住注意力，中间30秒展开剧情，最后12秒收尾和互动引导"
            ],
            "modification_memories": [],
            "general_memories": [
                "喜欢真实生活场景，贴近观众生活",
                "偏好自然表演，不做作",
                "避免过度夸张，保持真实感",
                "注重细节，让观众产生共鸣"
            ],
            "user_preferences": {
                "风格偏好": "轻松幽默，真实自然",
                "语气偏好": "自然随意，像朋友聊天",
                "平台偏好": "抖音",
                "内容偏好": "有趣有料，能引起共鸣"
            },
            "product_info": {},
            "shot_materials": [
                "镜头1: 开场悬念（吸引注意，制造好奇心，3秒抓住观众）",
                "镜头2: 日常场景（建立情境，生活化场景，让观众有代入感）",
                "镜头3: 意外情况发生（转折点，剧情推进，制造冲突）",
                "镜头4: 反应镜头（情绪表达，表情动作，夸张但不失真）",
                "镜头5: 反转或笑点（高潮，制造惊喜，让观众笑出声）",
                "镜头6: 结尾互动（引导评论，增加互动性，留下记忆点）"
            ]
        },
        {
            "id": "scenario_4",
            "name": "电商-服装-淘宝-复杂",
            "video_type": "ecommerce",
            "platform": "douyin",
            "duration": 60,  ***REMOVED*** 至少30秒
            "task_memories": [
                "推广春季气质连衣裙，面向25-35岁职场女性，这是一款适合春夏季节穿着的优雅知性连衣裙",
                "核心卖点：面料舒适透气适合春夏穿着，版型修身显瘦显高，多场景适用（职场通勤、约会聚会、休闲出游），性价比高适合日常穿着",
                "目标受众：25-35岁女性，追求优雅知性风格，注重品质和性价比，希望找到既实用又美观的服装",
                "风格要求：优雅知性，专业推荐但不生硬，真实展示不做过度修图，让消费者看到产品的真实状态",
                "需要详细展示：面料质感（材质手感、透气性说明）、上身效果（正面、侧面、背面多角度展示）、多场景穿搭（职场、约会、休闲等场景演示）、细节做工（缝线、纽扣、拉链等细节）、身材包容性（不同身材的穿着效果）、搭配建议（配饰、鞋包等整体造型）",
                "强调产品优势：面料舒适透气（优质棉质混纺，适合春夏穿着不闷热）、版型修身显瘦（剪裁设计显瘦显高，适合不同身材）、多场景适用（职场、约会、休闲都能穿，一衣多穿）、性价比高（200-300元价格区间，品质保证）",
                "内容要求：详细展示产品细节，真实上身效果不做过度修图，让消费者充分了解产品，有信心购买和推荐",
                "互动要求：引导观众分享自己的穿搭心得和搭配技巧，增加互动性和转化率，鼓励评论和提问"
            ],
            "modification_memories": [],
            "general_memories": [
                "喜欢详细展示细节，让消费者充分了解产品",
                "偏好真实上身效果，不做过度修图",
                "避免过度修图，保持真实感",
                "注重多角度展示，让消费者看到全貌"
            ],
            "user_preferences": {
                "风格偏好": "优雅知性，专业推荐",
                "语气偏好": "专业但不生硬，像朋友推荐",
                "平台偏好": "淘宝",
                "内容偏好": "详细实用，有参考价值"
            },
            "product_info": {
                "产品名称": "春季气质连衣裙",
                "核心卖点": "面料舒适透气，版型修身显瘦，多场景适用",
                "价格区间": "200-300元",
                "颜色": "米白色、浅蓝色、淡粉色",
                "尺码": "S-XL（适合不同身材）",
                "面料": "优质棉质混纺，透气舒适",
                "适用场景": "职场通勤、约会聚会、休闲出游"
            },
            "shot_materials": [
                "镜头1: 产品整体展示（连衣裙全貌，版型展示，优雅气质）",
                "镜头2: 面料细节特写（材质质感，手感展示，透气性说明）",
                "镜头3: 上身效果展示（正面、侧面、背面多角度上身效果）",
                "镜头4: 多场景穿搭演示（职场通勤、约会聚会、休闲出游等场景）",
                "镜头5: 细节做工展示（缝线、纽扣、拉链等细节，品质保证）",
                "镜头6: 身材包容性展示（不同身材的穿着效果，S-XL尺码展示）",
                "镜头7: 搭配建议（配饰、鞋包等搭配演示，整体造型建议）",
                "镜头8: 颜色选择建议（米白色、浅蓝色、淡粉色，不同肤色适配）"
            ]
        },
        {
            "id": "scenario_5",
            "name": "知识-科技-抖音-复杂",
            "video_type": "knowledge",
            "platform": "douyin",
            "duration": 90,  ***REMOVED*** 至少30秒
            "task_memories": [
                "AI技术科普，讲解GPT模型原理，面向对AI感兴趣的用户，帮助非专业技术人员理解AI技术",
                "核心内容：详细讲解GPT模型的工作原理（如何理解和生成文本）、训练过程（如何学习数据）、应用场景（聊天机器人、代码生成、文案创作等）、未来发展趋势（AI的未来发展方向）",
                "目标受众：对AI感兴趣的用户，有一定科技认知但非专业技术人员，希望通过视频了解AI技术的基本原理",
                "风格要求：专业但不晦涩，深入浅出，用类比解释复杂概念，让非专业观众也能理解AI技术",
                "需要平衡专业性和通俗性，避免过于技术化的术语，用生活化的例子和类比解释复杂概念，让观众在轻松的氛围中学习",
                "教学重点：什么是GPT（用生活例子解释）、GPT如何工作（工作原理的可视化展示）、GPT的应用场景（实际应用案例展示）、GPT的未来发展（引发思考的互动环节）",
                "内容要求：开头用生动例子引入降低理解门槛（比如把GPT比作什么），核心概念用类比解释（用通俗语言解释复杂概念），增加实际应用案例（贴近生活的AI应用），结尾引导思考AI未来发展（增加互动性）",
                "互动要求：引导观众思考AI的未来发展和影响，增加互动性和参与感，鼓励提问和讨论，建立学习社区氛围"
            ],
            "modification_memories": [],
            "general_memories": [
                "喜欢用类比解释复杂概念，让非专业观众也能理解",
                "偏好实际应用案例，让理论更贴近生活",
                "避免过于抽象的理论，注重实际应用",
                "强调AI的实际价值，让观众看到AI的未来潜力"
            ],
            "user_preferences": {
                "风格偏好": "专业但易懂，深入浅出",
                "语气偏好": "像老师一样耐心，娓娓道来",
                "平台偏好": "抖音",
                "内容偏好": "有深度有广度，能引发思考"
            },
            "product_info": {},
            "shot_materials": [
                "镜头1: 标题和讲师介绍（AI技术科普，建立权威性）",
                "镜头2: 生动例子引入（用生活例子解释AI，降低理解门槛）",
                "镜头3: 核心概念可视化（GPT模型原理图示，用图表展示）",
                "镜头4: 类比说明展示（用通俗语言解释复杂概念，比如把GPT比作什么）",
                "镜头5: 实际应用案例（贴近生活的AI应用：聊天机器人、代码生成、文案创作等）",
                "镜头6: 重点知识点标注（专业术语解释，关键概念强调）",
                "镜头7: 训练过程展示（GPT如何训练，数据如何学习）",
                "镜头8: 未来发展趋势（AI的未来发展方向，引发思考）",
                "镜头9: 互动环节（引导观众思考AI的未来，增加互动性）",
                "镜头10: 总结和回顾（本节重点回顾，引导观众继续学习）"
            ]
        }
    ]
    
    ***REMOVED*** 创建输出目录
    ***REMOVED*** 使用相对于脚本的路径，放在examples目录下
    script_dir = Path(__file__).parent
    output_dir = script_dir / "test_docs_new_template"
    output_dir.mkdir(exist_ok=True)
    
    print("=" * 70)
    print("基于新模板创建5份差异化的需求文档")
    print("=" * 70)
    print()
    
    created_files = []
    
    for scenario in scenarios:
        output_path = output_dir / f"{scenario['id']}.docx"
        try:
            create_diverse_doc(str(output_path), scenario)
            created_files.append(str(output_path))
            print(f"✅ 已创建: {scenario['name']}")
            print(f"   文件: {output_path}")
            print(f"   类型: {scenario['video_type']} | 平台: {scenario['platform']} | 时长: {scenario['duration']}秒")
            print(f"   镜头数: {len(scenario.get('shot_materials', []))}")
            print()
        except Exception as e:
            print(f"❌ 创建失败: {scenario['name']} - {e}")
            print()
    
    print("=" * 70)
    print(f"完成！共创建 {len(created_files)} 份文档")
    print(f"输出目录: {output_dir}")
    print("=" * 70)
    
    return created_files


if __name__ == "__main__":
    main()
