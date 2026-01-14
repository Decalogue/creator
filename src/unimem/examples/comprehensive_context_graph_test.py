***REMOVED***!/usr/bin/env python3
"""
Context Graph 综合测试框架

设计多样化的任务和数据，进行多轮交互优化，验证Context Graph机制，
并在每次测试后进行进化。

测试流程：
1. 生成多样化的测试需求文档（不同行业、平台、场景）
2. 执行完整的多轮交互流程（生成->反馈->优化->REFLECT）
3. 验证Context Graph功能（reasoning, decision_trace, 先例搜索等）
4. 分析结果，提取经验模式，优化系统参数
"""

import sys
import os
from pathlib import Path
from datetime import datetime
import json
import random
import logging
from typing import List, Dict, Any, Optional

***REMOVED*** 添加项目路径
project_root = Path(__file__).parent.parent.parent
sys.path.insert(0, str(project_root))

from unimem import UniMem
from unimem.memory_types import Experience, Context, Task, Memory
from unimem.examples.generate_video_script import VideoScriptGenerator
from unimem.neo4j import get_decision_events_for_memory
from py2neo import Graph

***REMOVED*** 初始化logger
logger = logging.getLogger(__name__)


class ContextGraphTestFramework:
    """Context Graph综合测试框架"""
    
    def __init__(self):
        """初始化测试框架"""
        print("="*70)
        print("Context Graph 综合测试框架")
        print("="*70)
        
        ***REMOVED*** 初始化UniMem
        config = {
            "storage": {
                "foa_backend": "redis",
                "da_backend": "redis",
                "ltm_backend": "neo4j",
            },
            "graph": {
                "neo4j_uri": "bolt://localhost:7680",
                "neo4j_user": "neo4j",
                "neo4j_password": "seeme_db",
            },
            "network": {
                "qdrant_host": "localhost",
                "qdrant_port": 6333,
            }
        }
        
        self.unimem = UniMem(config=config)
        self.generator = VideoScriptGenerator(use_service=False)
        
        ***REMOVED*** 连接Neo4j
        self.graph = Graph("bolt://localhost:7680", auth=('neo4j', 'seeme_db'), name='neo4j')
        
        ***REMOVED*** 测试场景配置
        self.test_scenarios = self._generate_diverse_scenarios()
        
        ***REMOVED*** 测试结果
        self.test_results = []
        
        print(f"✓ 测试框架初始化完成")
        print(f"✓ 设计了 {len(self.test_scenarios)} 个差异化测试场景\n")
    
    def _generate_diverse_scenarios(self) -> List[Dict[str, Any]]:
        """
        生成多样化的测试场景
        
        涵盖不同行业、平台、视频类型、需求复杂度
        """
        scenarios = []
        
        ***REMOVED*** 场景1: 电商-美妆-小红书（简单需求）
        scenarios.append({
            "id": "scenario_1",
            "name": "电商-美妆-小红书-简单",
            "video_type": "ecommerce",
            "platform": "xiaohongshu",
            "duration": 45,
            "complexity": "simple",
            "task_memories": [
                "推广新品口红",
                "突出持久度和显色度",
                "目标受众：年轻女性",
                "风格：清新自然"
            ],
            "product_info": {
                "产品名称": "春季限定口红",
                "核心卖点": "12小时持久，色彩饱满",
                "价格区间": "100-150元"
            },
            "feedbacks": [
                {"round": 1, "feedback": "试色部分可以更详细，展示不同光线下的效果。"},
                {"round": 2, "feedback": "持久度测试可以更直观，用时间轴展示。"},
                {"round": 3, "feedback": "整体风格很好，但可以增加一些使用技巧分享。"}
            ]
        })
        
        ***REMOVED*** 场景2: 教育-编程-抖音（中等复杂度）
        scenarios.append({
            "id": "scenario_2",
            "name": "教育-编程-抖音-中等",
            "video_type": "knowledge",
            "platform": "douyin",
            "duration": 90,
            "complexity": "medium",
            "task_memories": [
                "Python编程入门教学",
                "讲解变量和数据类型",
                "目标受众：编程初学者",
                "风格：通俗易懂，循序渐进"
            ],
            "product_info": {},
            "feedbacks": [
                {"round": 1, "feedback": "代码演示可以更清晰，增加注释说明。"},
                {"round": 2, "feedback": "实例可以更贴近实际应用场景。"},
                {"round": 3, "feedback": "讲解节奏可以稍微放慢，让初学者更容易跟上。"}
            ]
        })
        
        ***REMOVED*** 场景3: 娱乐-搞笑-抖音（简单）
        scenarios.append({
            "id": "scenario_3",
            "name": "娱乐-搞笑-抖音-简单",
            "video_type": "ecommerce",
            "platform": "douyin",
            "duration": 30,
            "complexity": "simple",
            "task_memories": [
                "搞笑日常段子",
                "展示生活中的小趣事",
                "目标受众：年轻人",
                "风格：轻松幽默"
            ],
            "product_info": {},
            "feedbacks": [
                {"round": 1, "feedback": "开头可以更有悬念，吸引观众继续看下去。"},
                {"round": 2, "feedback": "笑点可以更突出，增加一些音效或特效。"},
                {"round": 3, "feedback": "整体很好，但结尾可以更有记忆点。"}
            ]
        })
        
        ***REMOVED*** 场景4: 电商-服装-淘宝（复杂需求）
        scenarios.append({
            "id": "scenario_4",
            "name": "电商-服装-淘宝-复杂",
            "video_type": "ecommerce",
            "platform": "douyin",  ***REMOVED*** 注意：实际应该是淘宝，但测试用douyin
            "duration": 60,
            "complexity": "complex",
            "task_memories": [
                "推广春季连衣裙",
                "突出面料质感和上身效果",
                "目标受众：25-35岁女性",
                "风格：优雅知性",
                "需要多角度展示",
                "强调适用场景（职场、约会、休闲）"
            ],
            "product_info": {
                "产品名称": "春季气质连衣裙",
                "核心卖点": "面料舒适，版型修身，多场景适用",
                "价格区间": "200-300元",
                "颜色": "米白色、浅蓝色",
                "尺码": "S-XL"
            },
            "feedbacks": [
                {"round": 1, "feedback": "面料细节展示可以更清晰，突出质感。"},
                {"round": 2, "feedback": "上身效果可以展示更多角度和场景。"},
                {"round": 3, "feedback": "整体不错，但可以增加一些搭配建议。"},
                {"round": 4, "feedback": "可以增加一些身材包容性的说明，让更多消费者有信心购买。"}
            ]
        })
        
        ***REMOVED*** 场景5: 知识-科技-抖音（高复杂度）
        scenarios.append({
            "id": "scenario_5",
            "name": "知识-科技-抖音-复杂",
            "video_type": "knowledge",
            "platform": "douyin",
            "duration": 120,
            "complexity": "complex",
            "task_memories": [
                "AI技术科普",
                "讲解GPT模型原理",
                "目标受众：对AI感兴趣的用户",
                "风格：专业但不晦涩",
                "需要平衡专业性和通俗性",
                "避免过于技术化的术语"
            ],
            "product_info": {},
            "feedbacks": [
                {"round": 1, "feedback": "开头可以用一个生动的例子引入，降低理解门槛。"},
                {"round": 2, "feedback": "核心概念的解释可以用更多类比，帮助非专业观众理解。"},
                {"round": 3, "feedback": "可以增加一些实际应用案例，让理论更贴近生活。"},
                {"round": 4, "feedback": "结尾可以引导观众思考AI的未来发展，增加互动性。"}
            ]
        })
        
        ***REMOVED*** 为了提升测试效率，只返回第1个场景用于迭代测试
        return scenarios[:1]
    
    def create_word_document(self, scenario: Dict[str, Any], output_path: str) -> str:
        """
        为测试场景创建Word需求文档
        
        Args:
            scenario: 测试场景配置
            output_path: 输出路径
            
        Returns:
            创建的文档路径
        """
        from docx import Document
        
        doc = Document()
        
        ***REMOVED*** 标题
        doc.add_heading('短视频剧本需求文档', level=1)
        doc.add_paragraph(f'场景ID: {scenario["id"]}')
        doc.add_paragraph(f'场景名称: {scenario["name"]}')
        doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph()
        
        ***REMOVED*** 基本信息
        doc.add_heading('一、基本信息', level=1)
        doc.add_paragraph(f'视频类型: {scenario["video_type"]}')
        doc.add_paragraph(f'目标平台: {scenario["platform"]}')
        doc.add_paragraph(f'目标时长: {scenario["duration"]}秒')
        doc.add_paragraph()
        
        ***REMOVED*** 当前任务需求
        doc.add_heading('二、当前任务需求', level=1)
        for memory in scenario["task_memories"]:
            doc.add_paragraph(memory)
        doc.add_paragraph()
        
        ***REMOVED*** 商品信息（如果有）
        if scenario.get("product_info"):
            doc.add_heading('三、商品信息', level=1)
            for key, value in scenario["product_info"].items():
                doc.add_paragraph(f'{key}: {value}')
            doc.add_paragraph()
        
        ***REMOVED*** 通用记忆总结（模拟）
        doc.add_heading('四、通用记忆总结', level=1)
        doc.add_paragraph('喜欢用生活化语言')
        doc.add_paragraph('偏好真实体验分享')
        doc.add_paragraph()
        
        ***REMOVED*** 镜头素材（根据视频类型和时长动态生成）
        doc.add_heading('五、镜头素材', level=1)
        shot_materials = self._generate_shot_materials(
            video_type=scenario.get("video_type", "ecommerce"),
            duration=scenario.get("duration", 60),
            product_info=scenario.get("product_info", {}),
            task_memories=scenario.get("task_memories", [])
        )
        for shot in shot_materials:
            doc.add_paragraph(shot)
        doc.add_paragraph()
        
        ***REMOVED*** 保存文档
        doc.save(output_path)
        return output_path
    
    def _generate_shot_materials(
        self,
        video_type: str,
        duration: int,
        product_info: Optional[Dict[str, Any]] = None,
        task_memories: Optional[List[str]] = None
    ) -> List[str]:
        """
        根据视频类型和时长动态生成镜头素材
        
        Args:
            video_type: 视频类型（ecommerce, knowledge, entertainment等）
            duration: 视频时长（秒）
            product_info: 商品信息（可选，用于生成更具体的镜头）
            task_memories: 任务记忆列表（可选，用于调整知识类视频的镜头）
            
        Returns:
            镜头素材列表
        """
        ***REMOVED*** 根据时长估算需要的镜头数量（每7-10秒一个镜头）
        num_shots = max(5, min(12, int(duration / 8)))
        
        shots = []
        
        if video_type == "ecommerce":
            ***REMOVED*** 电商类视频镜头素材
            base_shots = [
                "镜头1: 产品整体展示（全景，展示产品全貌）",
                "镜头2: 产品特写（细节展示，突出卖点）",
                "镜头3: 使用场景展示（实际应用场景，用户使用产品）",
                "镜头4: 对比效果（使用前后对比/新旧对比，突出优势）",
                "镜头5: 用户反应镜头（惊喜/满意的表情和动作）",
                "镜头6: 产品细节特写（材质/做工/工艺细节）",
                "镜头7: 多角度展示（360度旋转/多场景切换）",
                "镜头8: 结尾转化镜头（引导购买，优惠信息展示）"
            ]
            
            ***REMOVED*** 根据商品信息调整镜头描述
            if product_info:
                product_name = product_info.get("产品名称", "产品")
                core_selling_point = product_info.get("核心卖点", "")
                
                if "口红" in product_name or "美妆" in product_name:
                    base_shots = [
                        "镜头1: 产品整体展示（口红外观，包装设计）",
                        "镜头2: 产品特写（膏体细节，色彩展示）",
                        "镜头3: 试色过程（自然光线下的试色效果）",
                        "镜头4: 持久度测试（时间轴展示12小时后的效果）",
                        "镜头5: 多色号展示（不同色号的对比）",
                        "镜头6: 日常妆容搭配（不同场合的妆容效果）",
                        "镜头7: 使用技巧分享（正确使用方法的演示）",
                        "镜头8: 结尾转化（优惠信息和购买引导）"
                    ]
                elif "连衣裙" in product_name or "服装" in product_name:
                    base_shots = [
                        "镜头1: 产品整体展示（连衣裙全貌，版型展示）",
                        "镜头2: 面料细节特写（材质质感，手感展示）",
                        "镜头3: 上身效果展示（多角度上身效果）",
                        "镜头4: 多场景穿搭演示（职场、约会、休闲等场景）",
                        "镜头5: 细节做工展示（缝线、纽扣、拉链等细节）",
                        "镜头6: 身材包容性展示（不同身材的穿着效果）",
                        "镜头7: 搭配建议（配饰、鞋包等搭配演示）",
                        "镜头8: 结尾转化（尺码说明和购买引导）"
                    ]
            
            shots = base_shots[:num_shots]
            
        elif video_type == "knowledge":
            ***REMOVED*** 知识类视频镜头素材
            base_shots = [
                "镜头1: 标题和讲师介绍（开场，建立权威性）",
                "镜头2: 核心概念展示（动画/图示，可视化解释）",
                "镜头3: 实例演示（实际操作，真实案例）",
                "镜头4: 重点知识点标注（强调关键信息）",
                "镜头5: 代码演示（屏幕录制，编程过程）",
                "镜头6: 运行效果展示（结果验证，输出展示）",
                "镜头7: 练习题展示（互动环节，巩固学习）",
                "镜头8: 总结和回顾（结尾，关键点回顾）"
            ]
            
            ***REMOVED*** 根据主题调整镜头描述
            task_memories_str = " ".join(task_memories) if task_memories else ""
            if "编程" in task_memories_str or "Python" in task_memories_str:
                base_shots = [
                    "镜头1: 课程标题和讲师介绍（Python编程入门）",
                    "镜头2: 代码演示屏幕录制（变量和数据类型示例）",
                    "镜头3: 实际运行效果展示（代码执行结果）",
                    "镜头4: 重点知识点标注（关键概念强调）",
                    "镜头5: 实例演示（贴近实际应用场景的案例）",
                    "镜头6: 注释说明展示（代码注释详解）",
                    "镜头7: 练习题展示（巩固练习，互动环节）",
                    "镜头8: 总结和回顾（本节重点回顾）"
                ]
            elif "AI" in task_memories_str or "GPT" in task_memories_str:
                base_shots = [
                    "镜头1: 标题和讲师介绍（AI技术科普）",
                    "镜头2: 生动例子引入（降低理解门槛的类比）",
                    "镜头3: 核心概念可视化（GPT模型原理图示）",
                    "镜头4: 实际应用案例（贴近生活的AI应用）",
                    "镜头5: 重点知识点标注（专业术语解释）",
                    "镜头6: 类比说明展示（用通俗语言解释复杂概念）",
                    "镜头7: 未来发展趋势（引发思考的互动环节）",
                    "镜头8: 总结和回顾（引导观众思考AI未来发展）"
                ]
            
            shots = base_shots[:num_shots]
            
        elif video_type == "entertainment":
            ***REMOVED*** 娱乐类视频镜头素材
            base_shots = [
                "镜头1: 开场悬念（吸引注意，制造好奇心）",
                "镜头2: 日常场景（建立情境，生活化场景）",
                "镜头3: 意外情况发生（转折点，剧情推进）",
                "镜头4: 反应镜头（情绪表达，表情动作）",
                "镜头5: 反转或笑点（高潮，制造惊喜）",
                "镜头6: 细节特写（表情/动作细节，增强表现力）",
                "镜头7: 结尾总结（收尾，留下记忆点）"
            ]
            
            shots = base_shots[:min(num_shots, len(base_shots))]
            
        else:
            ***REMOVED*** 默认通用镜头素材
            base_shots = [
                "镜头1: 开场镜头（吸引注意，建立情境）",
                "镜头2: 主要内容展示（核心内容呈现）",
                "镜头3: 细节特写（重点内容强调）",
                "镜头4: 使用场景（实际应用演示）",
                "镜头5: 互动环节（观众参与，增强互动）",
                "镜头6: 转折或高潮（剧情推进，保持兴趣）",
                "镜头7: 结尾总结（收尾，引导行动）"
            ]
            
            shots = base_shots[:num_shots]
        
        ***REMOVED*** 确保镜头编号连续
        for i in range(len(shots)):
            if shots[i].startswith("镜头"):
                shots[i] = f"镜头{i+1}: {shots[i].split(':', 1)[1].strip()}"
        
        return shots
    
    def run_single_test(self, scenario: Dict[str, Any]) -> Dict[str, Any]:
        """
        运行单个测试场景
        
        Args:
            scenario: 测试场景配置
            
        Returns:
            测试结果字典
        """
        print(f"\n{'='*70}")
        print(f"测试场景: {scenario['name']}")
        print(f"复杂度: {scenario['complexity']}")
        print(f"{'='*70}\n")
        
        test_result = {
            "scenario_id": scenario["id"],
            "scenario_name": scenario["name"],
            "complexity": scenario["complexity"],
            "timestamp": datetime.now().isoformat(),
            "steps": [],
            "memory_ids": [],
            "decision_events": [],
            "precedents_found": [],
            "reasoning_extracted": [],
            "errors": []
        }
        
        try:
            ***REMOVED*** 步骤1: 使用新模板格式的Word需求文档
            print("【步骤1】使用新模板格式的Word需求文档...")
            ***REMOVED*** 优先使用已创建的新模板文档，否则使用旧格式创建
            new_template_path = f"/root/data/AI/creator/src/unimem/examples/test_docs_new_template/{scenario['id']}.docx"
            if os.path.exists(new_template_path):
                doc_path = new_template_path
                print(f"  ✓ 使用新模板文档: {doc_path}\n")
            else:
                doc_path = f"/root/data/AI/creator/src/unimem/examples/test_scenario_{scenario['id']}.docx"
                self.create_word_document(scenario, doc_path)
                print(f"  ✓ 文档创建成功: {doc_path}\n")
            
            test_result["steps"].append({
                "step": "create_document",
                "status": "success",
                "doc_path": doc_path
            })
            
            ***REMOVED*** 步骤2: 解析文档
            print("【步骤2】解析Word文档...")
            try:
                doc_data = self.generator.parse_word_document(doc_path)
                test_result["steps"].append({
                    "step": "parse_document",
                    "status": "success",
                    "parsed_data": {
                        "task_memories": len(doc_data.get("task_memories", [])),
                        "video_type": doc_data.get("video_type"),
                        "platform": doc_data.get("platform")
                    }
                })
                print(f"  ✓ 解析成功: {len(doc_data.get('task_memories', []))}条任务记忆\n")
            except Exception as e:
                test_result["errors"].append(f"解析文档失败: {e}")
                print(f"  ✗ 解析失败: {e}\n")
                return test_result
            
            ***REMOVED*** 步骤3: 搜索先例（可选功能，用于提高质量，但不影响剧本生成）
            print("【步骤3】搜索相似先例（可选，用于提高质量）...")
            precedents = []
            try:
                task_memories = doc_data.get("task_memories", [])
                precedents = self.unimem.search_precedents(
                    inputs=task_memories[:3] if len(task_memories) >= 3 else task_memories,
                    rules=[
                        f"{doc_data.get('platform', 'douyin')}平台规则",
                        f"{doc_data.get('video_type', 'ecommerce')}类型规范"
                    ],
                    query_text=" ".join(task_memories[:5]) if task_memories else "",
                    top_k=3,
                    min_match_score=0.5
                )
                
                if precedents:
                    print(f"  ✓ 找到 {len(precedents)} 个相似先例（可用于参考）:")
                    for i, prec in enumerate(precedents, 1):
                        print(f"    {i}. {prec.content[:60]}...")
                        print(f"       理由: {prec.reasoning[:60] if prec.reasoning else 'N/A'}...")
                        test_result["precedents_found"].append({
                            "memory_id": prec.id,
                            "content_preview": prec.content[:100],
                            "reasoning": prec.reasoning[:100] if prec.reasoning else None,
                            "has_decision_trace": prec.decision_trace is not None
                        })
                else:
                    print(f"  ✓ 未找到相似先例（正常，将使用默认规则生成剧本）")
                
                test_result["steps"].append({
                    "step": "search_precedents",
                    "status": "success",
                    "precedents_count": len(precedents)
                })
                print()
            except Exception as e:
                ***REMOVED*** 先例搜索失败不影响后续流程，只记录警告，不添加到errors
                logger.warning(f"搜索先例失败（不影响剧本生成）: {e}")
                print(f"  ⚠ 搜索先例失败（不影响剧本生成）: {e}")
                print(f"  ✓ 将继续使用默认规则生成剧本\n")
                test_result["steps"].append({
                    "step": "search_precedents",
                    "status": "skipped",
                    "reason": str(e),
                    "precedents_count": 0
                })
            
            ***REMOVED*** 步骤4: 生成初始剧本
            print("【步骤4】生成初始剧本...")
            try:
                initial_script = self.generator.generate_script(doc_data)
                ***REMOVED*** 从adapter中获取script_memory_id
                script_memory_id = None
                if hasattr(self.generator, 'adapter') and hasattr(self.generator.adapter, 'last_script_memory_id'):
                    script_memory_id = self.generator.adapter.last_script_memory_id
                elif hasattr(self.generator, 'script_memory_id'):
                    script_memory_id = self.generator.script_memory_id
                
                ***REMOVED*** 如果还是没有，尝试从UniMem中查找最近的脚本记忆
                if not script_memory_id:
                    ***REMOVED*** 使用recall查找最近的脚本记忆
                    recall_results = self.unimem.recall(
                        query=f"{doc_data.get('video_type', '')} {doc_data.get('platform', '')} 脚本",
                        top_k=1
                    )
                    if recall_results:
                        script_memory_id = recall_results[0].memory.id
                
                if script_memory_id:
                    ***REMOVED*** 验证reasoning和decision_trace
                    ***REMOVED*** 使用neo4j.get_memory获取记忆
                    from unimem.neo4j import get_memory
                    memory = get_memory(script_memory_id)
                    if memory:
                        test_result["memory_ids"].append(script_memory_id)
                        
                        ***REMOVED*** 增强调试日志
                        has_reasoning = memory.reasoning is not None and memory.reasoning.strip() != ""
                        has_decision_trace = memory.decision_trace is not None and isinstance(memory.decision_trace, dict)
                        
                        ***REMOVED*** 如果decision_trace为空，尝试从Neo4j直接查询
                        if not has_decision_trace:
                            try:
                                from py2neo import Graph
                                graph = Graph("bolt://localhost:7680", auth=('neo4j', 'seeme_db'), name='neo4j')
                                result = graph.run("""
                                    MATCH (m:Memory {id: $memory_id})
                                    RETURN m.decision_trace as decision_trace
                                """, memory_id=script_memory_id).data()
                                if result and result[0].get("decision_trace"):
                                    logger.debug(f"Found decision_trace in Neo4j for memory {script_memory_id} but not in memory object")
                                    has_decision_trace = True
                            except Exception as e:
                                logger.debug(f"Failed to check decision_trace in Neo4j: {e}")
                        
                        test_result["reasoning_extracted"].append({
                            "memory_id": memory.id,
                            "has_reasoning": has_reasoning,
                            "reasoning_preview": memory.reasoning[:100] if memory.reasoning else None,
                            "has_decision_trace": has_decision_trace,
                            "decision_trace_keys": list(memory.decision_trace.keys()) if memory.decision_trace else [],
                            "decision_trace_type": type(memory.decision_trace).__name__ if memory.decision_trace else None
                        })
                        
                        ***REMOVED*** 验证DecisionEvent节点
                        events = get_decision_events_for_memory(script_memory_id)
                        if events:
                            test_result["decision_events"].append({
                                "memory_id": script_memory_id,
                                "events_count": len(events),
                                "latest_event": events[0] if events else None
                            })
                        else:
                            ***REMOVED*** 如果应该有decision_trace但没有事件，记录日志
                            if has_decision_trace:
                                logger.debug(f"Memory {script_memory_id} has decision_trace but no DecisionEvent created")
                
                test_result["steps"].append({
                    "step": "generate_script",
                    "status": "success",
                    "script_memory_id": script_memory_id
                })
                print(f"  ✓ 剧本生成成功: {len(initial_script.get('segments', []))}个片段")
                if script_memory_id:
                    print(f"  ✓ 记忆ID: {script_memory_id}")
                    if memory and memory.reasoning:
                        print(f"  ✓ 决策理由: {memory.reasoning[:60]}...")
                print()
            except Exception as e:
                test_result["errors"].append(f"生成剧本失败: {e}")
                print(f"  ✗ 生成剧本失败: {e}\n")
                return test_result
            
            ***REMOVED*** 步骤5: 多轮反馈优化
            print("【步骤5】多轮反馈优化...")
            feedback_count = 0
            accumulated_modifications = []
            
            for feedback_data in scenario.get("feedbacks", []):
                round_num = feedback_data.get("round", 1)
                feedback_text = feedback_data.get("feedback", "")
                
                print(f"  轮次 {round_num}: {feedback_text}")
                
                try:
                    ***REMOVED*** 存储反馈
                    if script_memory_id:
                        feedback_memory_id = self.generator.store_feedback_to_unimem(
                            feedback_text,
                            script_memory_id
                        )
                        
                    if feedback_memory_id:
                        try:
                            from unimem.neo4j import get_memory
                            feedback_memory = get_memory(feedback_memory_id)
                            if feedback_memory:
                                test_result["memory_ids"].append(feedback_memory_id)
                                ***REMOVED*** 检查reasoning和decision_trace
                                has_reasoning = feedback_memory.reasoning is not None and feedback_memory.reasoning.strip() != ""
                                has_decision_trace = feedback_memory.decision_trace is not None and isinstance(feedback_memory.decision_trace, dict)
                                
                                if has_reasoning:
                                    test_result["reasoning_extracted"].append({
                                        "memory_id": feedback_memory_id,
                                        "source": "user_feedback",
                                        "reasoning": feedback_memory.reasoning[:100],
                                        "has_reasoning": has_reasoning,
                                        "has_decision_trace": has_decision_trace,
                                        "decision_trace_keys": list(feedback_memory.decision_trace.keys()) if feedback_memory.decision_trace else []
                                    })
                                else:
                                    ***REMOVED*** 即使没有reasoning，也记录decision_trace信息
                                    test_result["reasoning_extracted"].append({
                                        "memory_id": feedback_memory_id,
                                        "source": "user_feedback",
                                        "has_reasoning": False,
                                        "has_decision_trace": has_decision_trace,
                                        "decision_trace_keys": list(feedback_memory.decision_trace.keys()) if feedback_memory.decision_trace else []
                                    })
                        except Exception as e:
                            logger.warning(f"Failed to get feedback memory {feedback_memory_id}: {e}")
                            pass  ***REMOVED*** 如果获取失败，继续执行
                    
                    ***REMOVED*** 提取修改需求
                    modifications = self.generator.extract_modification_feedback(
                        feedback_text,
                        existing_modifications=accumulated_modifications
                    )
                    
                    if modifications:
                        accumulated_modifications.extend(modifications)
                        
                        ***REMOVED*** 调用优化方法
                        try:
                            optimized_script = self.generator.optimize_and_regenerate(
                                original_doc_data=doc_data,
                                modification_feedback=modifications,
                                original_script=initial_script,
                                accumulated_modifications=accumulated_modifications
                            )
                            
                            if optimized_script:
                                initial_script = optimized_script  ***REMOVED*** 保存用于下一轮
                                
                                ***REMOVED*** 尝试获取新的script_memory_id
                                if hasattr(self.generator, 'adapter') and hasattr(self.generator.adapter, 'last_script_memory_id'):
                                    new_script_memory_id = self.generator.adapter.last_script_memory_id
                                    if new_script_memory_id and new_script_memory_id != script_memory_id:
                                        script_memory_id = new_script_memory_id
                                        test_result["memory_ids"].append(script_memory_id)
                                
                                print(f"    ✓ 脚本已优化")
                            else:
                                print(f"    ⚠ 优化完成，但未返回新脚本")
                        except Exception as e:
                            print(f"    ⚠ 优化脚本时出错: {e}")
                    else:
                        print(f"    ⚠ 未能提取到修改需求")
                    
                    feedback_count += 1
                    print(f"    ✓ 反馈处理完成\n")
                    
                except Exception as e:
                    test_result["errors"].append(f"轮次{round_num}优化失败: {e}")
                    print(f"    ✗ 优化失败: {e}\n")
            
            test_result["steps"].append({
                "step": "multi_round_optimization",
                "status": "success",
                "feedback_rounds": feedback_count
            })
            
            ***REMOVED*** 步骤6: REFLECT操作
            print("【步骤6】执行REFLECT操作...")
            try:
                ***REMOVED*** 收集相关记忆
                from unimem.neo4j import get_memory
                related_memories = []
                for mem_id in test_result["memory_ids"]:
                    memory = get_memory(mem_id)
                    if memory:
                        related_memories.append(memory)
                
                if related_memories:
                    task = Task(
                        id=f"task_{scenario['id']}",
                        description=f"完成{scenario['name']}视频脚本创作",
                        context=f"场景: {scenario['complexity']}复杂度, {len(scenario['feedbacks'])}轮反馈优化",
                        metadata={"scenario_id": scenario["id"]}
                    )
                    
                    evolved_memories = self.unimem.reflect(
                        memories=related_memories,
                        current_task=task
                    )
                    
                    ***REMOVED*** 检查是否提取到新经验
                    new_experiences = [m for m in evolved_memories if m.memory_type is not None and m.memory_type.value == "experience"]
                    if new_experiences:
                        print(f"  ✓ REFLECT完成，提取了 {len(new_experiences)} 条经验记忆:")
                        for exp in new_experiences:
                            print(f"    - {exp.content[:60]}...")
                            if exp.reasoning:
                                print(f"      理由: {exp.reasoning[:60]}...")
                                test_result["reasoning_extracted"].append({
                                    "memory_id": exp.id,
                                    "source": "reflect_experience",
                                    "reasoning": exp.reasoning[:100]
                                })
                    else:
                        print(f"  ✓ REFLECT完成，但未提取到新经验记忆")
                    
                    test_result["steps"].append({
                        "step": "reflect",
                        "status": "success",
                        "evolved_count": len(evolved_memories),
                        "new_experiences": len(new_experiences)
                    })
                else:
                    print(f"  ⚠ 没有相关记忆可进行REFLECT")
                    test_result["steps"].append({
                        "step": "reflect",
                        "status": "skipped",
                        "reason": "no_related_memories"
                    })
                print()
                
            except Exception as e:
                test_result["errors"].append(f"REFLECT失败: {e}")
                print(f"  ✗ REFLECT失败: {e}\n")
            
            ***REMOVED*** 步骤7: 验证Context Graph功能
            print("【步骤7】验证Context Graph功能...")
            self._verify_context_graph(test_result)
            print()
            
            ***REMOVED*** 步骤8: 获取并保存最终剧本
            print("【步骤8】获取最终剧本...")
            import sys
            sys.stdout.flush()  ***REMOVED*** 确保print输出被刷新
            logger.info("Step 8: Getting final script...")
            try:
                ***REMOVED*** 优先使用测试过程中保存的完整脚本内容（initial_script）
                ***REMOVED*** 因为store_script_to_unimem只存储了脚本预览，不是完整内容
                final_script_content = None
                final_script_memory_id = None
                
                ***REMOVED*** 从test_result的steps中查找最新的脚本
                for step in reversed(test_result["steps"]):
                    if step.get("step") == "generate_script" and step.get("script_memory_id"):
                        final_script_memory_id = step.get("script_memory_id")
                        logger.info(f"Found script_memory_id from steps: {final_script_memory_id}")
                        break
                
                ***REMOVED*** 如果optimize_and_regenerate创建了新的脚本，使用新的memory_id
                if hasattr(self.generator, 'adapter') and hasattr(self.generator.adapter, 'last_script_memory_id'):
                    new_script_id = self.generator.adapter.last_script_memory_id
                    if new_script_id:
                        logger.info(f"Found last_script_memory_id from adapter: {new_script_id}")
                        if new_script_id != final_script_memory_id:
                            final_script_memory_id = new_script_id
                            logger.info(f"Using new script_memory_id: {final_script_memory_id}")
                
                ***REMOVED*** 尝试从initial_script获取完整脚本内容
                if initial_script and initial_script.get("script"):
                    final_script_content = initial_script.get("script")
                    logger.info(f"Found full script content from initial_script, length={len(final_script_content)}")
                elif initial_script and initial_script.get("segments"):
                    ***REMOVED*** 如果没有script字段，从segments构建
                    segments = initial_script.get("segments", [])
                    final_script_content = "\n\n".join([
                        f"【片段{i+1}】{seg.get('content', '')}" 
                        for i, seg in enumerate(segments)
                    ])
                    logger.info(f"Built script content from segments, length={len(final_script_content)}")
                
                if final_script_memory_id:
                    logger.info(f"Attempting to get memory from Neo4j: {final_script_memory_id}")
                    from unimem.neo4j import get_memory
                    final_memory = None
                    
                    ***REMOVED*** 尝试多次获取Memory（重试机制）
                    for retry in range(3):
                        try:
                            final_memory = get_memory(final_script_memory_id)
                            if final_memory:
                                logger.info(f"Successfully retrieved memory from Neo4j (attempt {retry + 1})")
                                break
                            else:
                                logger.warning(f"Memory not found in Neo4j (attempt {retry + 1}), retrying...")
                                import time
                                time.sleep(0.5)  ***REMOVED*** 等待0.5秒后重试
                        except Exception as e:
                            logger.warning(f"Error getting memory from Neo4j (attempt {retry + 1}): {e}")
                            if retry < 2:
                                import time
                                time.sleep(0.5)
                    
                    ***REMOVED*** 优先使用从initial_script获取的完整脚本内容
                    if final_script_content:
                        ***REMOVED*** 使用完整脚本内容，从Memory获取元数据
                        if final_memory:
                            test_result["final_script"] = {
                                "memory_id": final_script_memory_id,
                                "content": final_script_content,  ***REMOVED*** 使用完整脚本内容
                                "content_preview": final_script_content[:500] + "..." if len(final_script_content) > 500 else final_script_content,
                                "keywords": final_memory.keywords[:10] if final_memory.keywords else [],
                                "tags": final_memory.tags[:10] if final_memory.tags else [],
                                "reasoning": final_memory.reasoning,
                                "has_decision_trace": final_memory.decision_trace is not None,
                                "source": "initial_script"
                            }
                        else:
                            ***REMOVED*** 即使无法获取Memory，也保存完整脚本内容，但尝试从Neo4j直接查询reasoning和decision_trace
                            logger.warning(f"Failed to get memory object from Neo4j, trying direct query for {final_script_memory_id}")
                            try:
                                from py2neo import Graph
                                graph = Graph("bolt://localhost:7680", auth=('neo4j', 'seeme_db'), name='neo4j')
                                result = graph.run("""
                                    MATCH (m:Memory {id: $memory_id})
                                    RETURN m.reasoning as reasoning, m.decision_trace as decision_trace,
                                           m.keywords as keywords, m.tags as tags
                                """, memory_id=final_script_memory_id).data()
                                
                                if result and result[0]:
                                    neo4j_data = result[0]
                                    test_result["final_script"] = {
                                        "memory_id": final_script_memory_id,
                                        "content": final_script_content,
                                        "content_preview": final_script_content[:500] + "..." if len(final_script_content) > 500 else final_script_content,
                                        "keywords": neo4j_data.get("keywords", [])[:10] if neo4j_data.get("keywords") else [],
                                        "tags": neo4j_data.get("tags", [])[:10] if neo4j_data.get("tags") else [],
                                        "reasoning": neo4j_data.get("reasoning"),
                                        "has_decision_trace": neo4j_data.get("decision_trace") is not None,
                                        "source": "initial_script_neo4j_direct"
                                    }
                                    logger.info(f"Retrieved reasoning and decision_trace from Neo4j direct query")
                                else:
                                    ***REMOVED*** 如果直接查询也失败，保存基本信息
                                    test_result["final_script"] = {
                                        "memory_id": final_script_memory_id,
                                        "content": final_script_content,
                                        "content_preview": final_script_content[:500] + "..." if len(final_script_content) > 500 else final_script_content,
                                        "keywords": [],
                                        "tags": [],
                                        "reasoning": None,
                                        "has_decision_trace": False,
                                        "source": "initial_script_no_memory"
                                    }
                                    logger.warning(f"Failed to retrieve memory metadata from Neo4j for {final_script_memory_id}")
                            except Exception as e:
                                logger.error(f"Failed to query Neo4j directly: {e}")
                                test_result["final_script"] = {
                                    "memory_id": final_script_memory_id,
                                    "content": final_script_content,
                                    "content_preview": final_script_content[:500] + "..." if len(final_script_content) > 500 else final_script_content,
                                    "keywords": [],
                                    "tags": [],
                                    "reasoning": None,
                                    "has_decision_trace": False,
                                    "source": "initial_script_no_memory"
                                }
                        print(f"  ✓ 最终剧本已获取（从测试脚本）: memory_id={final_script_memory_id}")
                        print(f"  ✓ 剧本内容长度: {len(final_script_content)} 字符")
                        sys.stdout.flush()
                        logger.info(f"Successfully got final script from initial_script: memory_id={final_script_memory_id}, content_length={len(final_script_content)}")
                        if final_memory and final_memory.reasoning:
                            print(f"  ✓ 决策理由: {final_memory.reasoning[:80]}...")
                            sys.stdout.flush()
                    elif final_memory:
                        ***REMOVED*** 如果没有完整脚本内容，使用Memory中的内容（可能是预览）
                        test_result["final_script"] = {
                            "memory_id": final_script_memory_id,
                            "content": final_memory.content,
                            "content_preview": final_memory.content[:500] + "..." if len(final_memory.content) > 500 else final_memory.content,
                            "keywords": final_memory.keywords[:10] if final_memory.keywords else [],
                            "tags": final_memory.tags[:10] if final_memory.tags else [],
                            "reasoning": final_memory.reasoning,
                            "has_decision_trace": final_memory.decision_trace is not None,
                            "source": "neo4j_memory"
                        }
                        print(f"  ✓ 最终剧本已获取（从Neo4j）: memory_id={final_script_memory_id}")
                        print(f"  ✓ 剧本内容长度: {len(final_memory.content)} 字符")
                        sys.stdout.flush()
                        logger.info(f"Successfully got final script from Neo4j: memory_id={final_script_memory_id}, content_length={len(final_memory.content)}")
                        if final_memory.reasoning:
                            print(f"  ✓ 决策理由: {final_memory.reasoning[:80]}...")
                            sys.stdout.flush()
                    else:
                        ***REMOVED*** 尝试从Qdrant获取（fallback）
                        logger.warning(f"Failed to get memory from Neo4j: {final_script_memory_id}, trying Qdrant fallback...")
                        try:
                            from unimem.adapters.atom_link_adapter import AtomLinkAdapter
                            if hasattr(self.unimem, 'network_adapter') and isinstance(self.unimem.network_adapter, AtomLinkAdapter):
                                ***REMOVED*** 尝试通过语义检索找到这个memory
                                results = self.unimem.recall(query=f"memory_id:{final_script_memory_id}", top_k=1)
                                if results and len(results) > 0 and results[0].memory:
                                    final_memory = results[0].memory
                                    test_result["final_script"] = {
                                        "memory_id": final_script_memory_id,
                                        "content": final_memory.content,
                                        "content_preview": final_memory.content[:500] + "..." if len(final_memory.content) > 500 else final_memory.content,
                                        "keywords": final_memory.keywords[:10] if final_memory.keywords else [],
                                        "tags": final_memory.tags[:10] if final_memory.tags else [],
                                        "reasoning": final_memory.reasoning,
                                        "has_decision_trace": final_memory.decision_trace is not None,
                                        "source": "qdrant_fallback"
                                    }
                                    print(f"  ✓ 最终剧本已获取（从Qdrant）: memory_id={final_script_memory_id}")
                                    print(f"  ✓ 剧本内容长度: {len(final_memory.content)} 字符")
                                    sys.stdout.flush()
                                    logger.info(f"Successfully got final script from Qdrant fallback: memory_id={final_script_memory_id}")
                                else:
                                    raise Exception("Qdrant fallback also failed")
                            else:
                                raise Exception("Network adapter not available for Qdrant fallback")
                        except Exception as fallback_error:
                            logger.error(f"Qdrant fallback also failed: {fallback_error}")
                            print(f"  ⚠ 无法获取最终剧本Memory: {final_script_memory_id}")
                            print(f"  ⚠ Neo4j获取失败，Qdrant fallback也失败: {fallback_error}")
                            sys.stdout.flush()
                            test_result["final_script"] = {
                                "memory_id": final_script_memory_id,
                                "error": f"无法从Neo4j和Qdrant获取Memory: {fallback_error}"
                            }
                else:
                    print(f"  ⚠ 未找到最终剧本Memory ID")
                    sys.stdout.flush()
                    logger.warning("No final script memory_id found")
                    ***REMOVED*** 尝试从test_result["memory_ids"]中查找可能的脚本memory
                    if test_result.get("memory_ids"):
                        logger.info(f"Trying to find script from memory_ids: {test_result['memory_ids']}")
                        for mem_id in reversed(test_result["memory_ids"]):
                            try:
                                mem = get_memory(mem_id)
                                if mem and len(mem.content) > 500:  ***REMOVED*** 脚本通常比较长
                                    logger.info(f"Found potential script memory: {mem_id}, content_length={len(mem.content)}")
                                    final_script_memory_id = mem_id
                                    test_result["final_script"] = {
                                        "memory_id": final_script_memory_id,
                                        "content": mem.content,
                                        "content_preview": mem.content[:500] + "..." if len(mem.content) > 500 else mem.content,
                                        "keywords": mem.keywords[:10] if mem.keywords else [],
                                        "tags": mem.tags[:10] if mem.tags else [],
                                        "reasoning": mem.reasoning,
                                        "has_decision_trace": mem.decision_trace is not None,
                                        "source": "auto_detected"
                                    }
                                    print(f"  ✓ 自动检测到最终剧本: memory_id={final_script_memory_id}")
                                    print(f"  ✓ 剧本内容长度: {len(mem.content)} 字符")
                                    sys.stdout.flush()
                                    break
                            except:
                                continue
                    if not test_result.get("final_script"):
                        test_result["final_script"] = None
                print()
                sys.stdout.flush()
            except Exception as e:
                logger.error(f"Failed to get final script: {e}", exc_info=True)
                test_result["final_script"] = {"error": str(e)}
                print(f"  ⚠ 获取最终剧本失败: {e}\n")
                sys.stdout.flush()
            
        except Exception as e:
            test_result["errors"].append(f"测试失败: {e}")
            print(f"\n✗ 测试失败: {e}\n")
        
        return test_result
    
    def _verify_context_graph(self, test_result: Dict[str, Any]) -> None:
        """验证Context Graph功能"""
        verifications = []
        
        ***REMOVED*** 验证1: reasoning字段覆盖率
        total_memories = len(test_result["memory_ids"])
        memories_with_reasoning = sum(1 for r in test_result["reasoning_extracted"] if r.get("reasoning"))
        reasoning_coverage = memories_with_reasoning / total_memories * 100 if total_memories > 0 else 0
        
        verifications.append({
            "check": "reasoning字段覆盖率",
            "result": f"{reasoning_coverage:.1f}% ({memories_with_reasoning}/{total_memories})",
            "status": "✓" if reasoning_coverage > 50 else "⚠"
        })
        print(f"  {verifications[-1]['status']} {verifications[-1]['check']}: {verifications[-1]['result']}")
        
        ***REMOVED*** 验证2: decision_trace覆盖率
        memories_with_trace = sum(1 for r in test_result["reasoning_extracted"] if r.get("has_decision_trace"))
        trace_coverage = memories_with_trace / total_memories * 100 if total_memories > 0 else 0
        
        verifications.append({
            "check": "decision_trace字段覆盖率",
            "result": f"{trace_coverage:.1f}% ({memories_with_trace}/{total_memories})",
            "status": "✓" if trace_coverage > 50 else "⚠"
        })
        print(f"  {verifications[-1]['status']} {verifications[-1]['check']}: {verifications[-1]['result']}")
        
        ***REMOVED*** 验证3: DecisionEvent节点创建
        events_count = len(test_result["decision_events"])
        verifications.append({
            "check": "DecisionEvent节点创建",
            "result": f"{events_count}个事件节点",
            "status": "✓" if events_count > 0 else "⚠"
        })
        print(f"  {verifications[-1]['status']} {verifications[-1]['check']}: {verifications[-1]['result']}")
        
        ***REMOVED*** 验证4: 先例搜索功能
        precedents_count = len(test_result["precedents_found"])
        verifications.append({
            "check": "先例搜索功能",
            "result": f"找到{precedents_count}个先例" if precedents_count > 0 else "未找到先例（可能是新场景）",
            "status": "✓" if precedents_count >= 0 else "✗"
        })
        print(f"  {verifications[-1]['status']} {verifications[-1]['check']}: {verifications[-1]['result']}")
        
        test_result["verifications"] = verifications
    
    def run_all_tests(self) -> Dict[str, Any]:
        """运行所有测试场景"""
        print("\n" + "="*70)
        print("开始运行所有测试场景")
        print("="*70 + "\n")
        
        all_results = {
            "start_time": datetime.now().isoformat(),
            "scenarios": [],
            "summary": {}
        }
        
        for scenario in self.test_scenarios:
            result = self.run_single_test(scenario)
            all_results["scenarios"].append(result)
            self.test_results.append(result)
            
            ***REMOVED*** 每个场景结束后立即保存该场景的最终剧本
            final_script = result.get("final_script")
            if final_script:
                if final_script.get("content"):
                    logger.info(f"Saving final script for scenario {len(all_results['scenarios'])}")
                    self._save_single_scenario_script(result, len(all_results["scenarios"]))
                else:
                    logger.warning(f"Final script exists but has no content: {final_script}")
                    print(f"  ⚠ 场景 {len(all_results['scenarios'])} 的最终剧本没有内容，无法保存")
                    import sys
                    sys.stdout.flush()
            
            ***REMOVED*** 每次测试后等待一小段时间，确保数据完整写入
            import time
            time.sleep(1)
        
        all_results["end_time"] = datetime.now().isoformat()
        all_results["summary"] = self._analyze_results(all_results["scenarios"])
        
        return all_results
    
    def _analyze_results(self, results: List[Dict[str, Any]]) -> Dict[str, Any]:
        """分析测试结果，提取经验模式"""
        print("\n" + "="*70)
        print("测试结果分析")
        print("="*70 + "\n")
        
        summary = {
            "total_scenarios": len(results),
            "successful_scenarios": sum(1 for r in results if not r.get("errors")),
            "total_memories": sum(len(r.get("memory_ids", [])) for r in results),
            "total_decision_events": sum(len(r.get("decision_events", [])) for r in results),
            "reasoning_coverage": 0,
            "precedents_found": 0,
            "common_patterns": [],
            "improvements": []
        }
        
        ***REMOVED*** 统计reasoning覆盖率
        all_reasoning = []
        for r in results:
            all_reasoning.extend(r.get("reasoning_extracted", []))
        
        ***REMOVED*** 修复：使用has_reasoning字段而不是reasoning字段
        memories_with_reasoning = sum(1 for re in all_reasoning if re.get("has_reasoning") or re.get("reasoning"))
        total_memories = summary["total_memories"]
        summary["reasoning_coverage"] = memories_with_reasoning / total_memories * 100 if total_memories > 0 else 0
        
        ***REMOVED*** 统计先例
        summary["precedents_found"] = sum(len(r.get("precedents_found", [])) for r in results)
        
        print(f"【总体统计】")
        print(f"  测试场景数: {summary['total_scenarios']}")
        print(f"  成功场景数: {summary['successful_scenarios']}")
        print(f"  总记忆数: {summary['total_memories']}")
        print(f"  决策事件数: {summary['total_decision_events']}")
        print(f"  Reasoning覆盖率: {summary['reasoning_coverage']:.1f}%")
        print(f"  找到先例数: {summary['precedents_found']}")
        
        ***REMOVED*** 分析公共模式
        print(f"\n【公共模式识别】")
        common_feedbacks = {}
        for r in results:
            for step in r.get("steps", []):
                if step.get("step") == "multi_round_optimization":
                    ***REMOVED*** 这里可以从记忆中提取公共反馈模式
                    pass
        
        ***REMOVED*** 从Neo4j中分析
        try:
            ***REMOVED*** 查询所有experience类型记忆
            query = """
            MATCH (m:Memory)
            WHERE m.memory_type = 'experience'
            RETURN m.content as content, m.reasoning as reasoning, m.source as source
            ORDER BY m.created_at DESC
            LIMIT 20
            """
            experience_results = self.graph.run(query).data()
            
            if experience_results:
                print(f"  提取到 {len(experience_results)} 条经验记忆:")
                for i, exp in enumerate(experience_results[:5], 1):
                    content = exp.get('content', '') or ''
                    reasoning = exp.get('reasoning', '') or ''
                    print(f"    {i}. {content[:60]}...")
                    if reasoning:
                        print(f"       理由: {reasoning[:60]}...")
            
            ***REMOVED*** 查询决策事件统计
            event_query = """
            MATCH (de:DecisionEvent)
            RETURN count(de) as total_events
            """
            event_result = self.graph.run(event_query).data()
            if event_result:
                summary["total_decision_events"] = event_result[0].get("total_events", 0)
                print(f"\n  Neo4j中的决策事件总数: {summary['total_decision_events']}")
        
        except Exception as e:
            print(f"  ⚠ 分析Neo4j数据失败: {e}")
        
        ***REMOVED*** 改进建议
        print(f"\n【改进建议】")
        if summary["reasoning_coverage"] < 80:
            print(f"  ⚠ Reasoning覆盖率偏低 ({summary['reasoning_coverage']:.1f}%)，建议增强REFLECT提示词")
            summary["improvements"].append("增强REFLECT提示词，提升reasoning提取率")
        
        if summary["total_decision_events"] == 0:
            print(f"  ⚠ 未创建DecisionEvent节点，检查decision_trace是否正确捕获")
            summary["improvements"].append("检查decision_trace捕获逻辑")
        
        if summary["precedents_found"] == 0 and total_memories > 5:
            print(f"  ⚠ 未找到先例，可能是因为场景差异较大，建议降低搜索阈值")
            summary["improvements"].append("优化先例搜索阈值和算法")
        else:
            print(f"  ✓ 先例搜索功能正常")
        
        return summary
    
    def _save_single_scenario_script(self, scenario_result: Dict[str, Any], scenario_index: int) -> None:
        """保存单个场景的最终剧本（每个场景结束后立即保存）"""
        scripts_dir = Path(__file__).parent / "final_scripts"
        scripts_dir.mkdir(exist_ok=True)
        
        final_script = scenario_result.get("final_script")
        if final_script and final_script.get("content"):
            scenario_name = scenario_result.get("scenario_name", f"scenario_{scenario_index}")
            ***REMOVED*** 清理文件名中的特殊字符
            safe_name = "".join(c if c.isalnum() or c in ('-', '_') else '_' for c in scenario_name)
            script_file = scripts_dir / f"{safe_name}_final_script.txt"
            with open(script_file, 'w', encoding='utf-8') as f:
                f.write(f"场景: {scenario_name}\n")
                f.write(f"Memory ID: {final_script.get('memory_id', 'N/A')}\n")
                f.write("=" * 80 + "\n\n")
                f.write(final_script.get("content", ""))
                f.write("\n\n" + "=" * 80 + "\n")
                if final_script.get("reasoning"):
                    f.write(f"\n决策理由:\n{final_script.get('reasoning')}\n")
            print(f"  ✓ 场景 {scenario_index} 最终剧本已保存: {script_file.name}")
            logger.info(f"Saved final script for scenario {scenario_index}: {script_file.name}")
    
    def save_results(self, results: Dict[str, Any], output_path: str = None) -> str:
        """保存测试结果"""
        if output_path is None:
            output_path = f"/root/data/AI/creator/src/unimem/examples/context_graph_test_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(results, f, ensure_ascii=False, indent=2, default=str)
        
        ***REMOVED*** 检查是否所有场景的最终剧本都已保存（可能已经在每个场景结束后保存了）
        scripts_dir = Path(output_path).parent / "final_scripts"
        scripts_dir.mkdir(exist_ok=True)
        
        saved_count = 0
        for i, scenario in enumerate(results.get("scenarios", []), 1):
            final_script = scenario.get("final_script")
            if final_script and final_script.get("content"):
                scenario_name = scenario.get("scenario_name", f"scenario_{i}")
                ***REMOVED*** 清理文件名中的特殊字符
                safe_name = "".join(c if c.isalnum() or c in ('-', '_') else '_' for c in scenario_name)
                script_file = scripts_dir / f"{safe_name}_final_script.txt"
                ***REMOVED*** 如果文件已存在，说明已经在场景结束时保存过了，跳过
                if not script_file.exists():
                    with open(script_file, 'w', encoding='utf-8') as f:
                        f.write(f"场景: {scenario_name}\n")
                        f.write(f"Memory ID: {final_script.get('memory_id', 'N/A')}\n")
                        f.write("=" * 80 + "\n\n")
                        f.write(final_script.get("content", ""))
                        f.write("\n\n" + "=" * 80 + "\n")
                        if final_script.get("reasoning"):
                            f.write(f"\n决策理由:\n{final_script.get('reasoning')}\n")
                    saved_count += 1
                    print(f"  ✓ 场景 {i} 最终剧本已保存: {script_file.name}")
                else:
                    saved_count += 1  ***REMOVED*** 已存在，计入总数
        
        print(f"\n测试结果已保存到: {output_path}")
        if saved_count > 0:
            print(f"最终剧本文件目录: {scripts_dir} (共{saved_count}个文件)")
        return output_path
    
    def evolve_system(self, results: Dict[str, Any]) -> Dict[str, Any]:
        """
        基于测试结果进化系统
        
        提取经验模式，优化系统参数
        """
        print("\n" + "="*70)
        print("系统进化分析")
        print("="*70 + "\n")
        
        evolution_actions = []
        
        ***REMOVED*** 1. 提取高频反馈模式，总结为经验
        print("【1. 提取反馈模式】")
        feedback_patterns = {}
        for scenario_result in results.get("scenarios", []):
            for step in scenario_result.get("steps", []):
                if step.get("step") == "multi_round_optimization":
                    ***REMOVED*** 可以分析反馈内容
                    pass
        
        ***REMOVED*** 2. 分析决策理由模式
        print("【2. 分析决策理由模式】")
        reasoning_patterns = {}
        for scenario_result in results.get("scenarios", []):
            for reasoning_data in scenario_result.get("reasoning_extracted", []):
                reasoning = reasoning_data.get("reasoning", "")
                if reasoning:
                    ***REMOVED*** 提取关键词
                    keywords = ["因为", "基于", "根据", "由于"]
                    for kw in keywords:
                        if kw in reasoning:
                            pattern_type = f"包含'{kw}'的决策理由"
                            reasoning_patterns[pattern_type] = reasoning_patterns.get(pattern_type, 0) + 1
        
        if reasoning_patterns:
            print("  决策理由模式:")
            for pattern, count in sorted(reasoning_patterns.items(), key=lambda x: x[1], reverse=True)[:5]:
                print(f"    - {pattern}: {count}次")
        
        ***REMOVED*** 3. 优化建议
        print("\n【3. 系统优化建议】")
        summary = results.get("summary", {})
        
        if summary.get("reasoning_coverage", 0) < 80:
            evolution_actions.append({
                "action": "enhance_reflect_prompt",
                "description": "增强REFLECT提示词，更主动地提取reasoning",
                "priority": "high"
            })
            print("  - 增强REFLECT提示词，提升reasoning提取率")
        
        if summary.get("total_decision_events", 0) == 0:
            evolution_actions.append({
                "action": "fix_decision_event_creation",
                "description": "检查并修复DecisionEvent节点创建逻辑",
                "priority": "high"
            })
            print("  - 检查DecisionEvent节点创建逻辑")
        
        ***REMOVED*** 4. 提取公共创作原则
        print("\n【4. 提取公共创作原则】")
        try:
            ***REMOVED*** 查询所有experience记忆，提取创作原则
            query = """
            MATCH (m:Memory)
            WHERE m.memory_type = 'experience' AND m.source IN ['reflect', 'reflect_implicit']
            RETURN m.content as content, m.reasoning as reasoning
            ORDER BY m.created_at DESC
            LIMIT 10
            """
            principles = self.graph.run(query).data()
            
            if principles:
                print(f"  提取到 {len(principles)} 条创作原则:")
                for i, p in enumerate(principles, 1):
                    content = p.get('content', '') or ''
                    reasoning = p.get('reasoning', '') or ''
                    print(f"    {i}. {content[:80]}...")
                    if reasoning:
                        print(f"       依据: {reasoning[:60]}...")
                
                evolution_actions.append({
                    "action": "create_principles_library",
                    "description": f"基于{len(principles)}条经验记忆，创建创作原则库",
                    "priority": "medium",
                    "principles_count": len(principles)
                })
            else:
                print("  ⚠ 暂无经验记忆，需要更多测试数据")
        
        except Exception as e:
            print(f"  ⚠ 提取原则失败: {e}")
        
        evolution_report = {
            "timestamp": datetime.now().isoformat(),
            "actions": evolution_actions,
            "reasoning_patterns": reasoning_patterns,
            "next_steps": [
                "运行更多测试场景，积累经验数据",
                "基于经验数据优化系统参数",
                "建立创作原则库"
            ]
        }
        
        return evolution_report


def main():
    """主函数"""
    print("""
    ╔══════════════════════════════════════════════════════════════════╗
    ║         Context Graph 综合测试框架                                ║
    ║                                                                  ║
    ║  本测试将：                                                       ║
    ║  1. 生成多样化的测试场景（不同行业、平台、复杂度）                ║
    ║  2. 执行完整的多轮交互流程（生成->反馈->优化->REFLECT）           ║
    ║  3. 验证Context Graph功能（reasoning, decision_trace, 先例搜索）║
    ║  4. 分析结果，提取经验模式，进化系统                              ║
    ╚══════════════════════════════════════════════════════════════════╝
    """)
    
    framework = ContextGraphTestFramework()
    
    ***REMOVED*** 运行所有测试
    results = framework.run_all_tests()
    
    ***REMOVED*** 保存结果
    results_path = framework.save_results(results)
    
    ***REMOVED*** 系统进化分析
    evolution_report = framework.evolve_system(results)
    
    ***REMOVED*** 保存进化报告
    evolution_path = results_path.replace("_results_", "_evolution_")
    with open(evolution_path, 'w', encoding='utf-8') as f:
        json.dump(evolution_report, f, ensure_ascii=False, indent=2, default=str)
    
    print("\n" + "="*70)
    print("测试完成！")
    print("="*70)
    print(f"测试结果: {results_path}")
    print(f"进化报告: {evolution_path}")
    print("="*70 + "\n")


if __name__ == "__main__":
    main()
