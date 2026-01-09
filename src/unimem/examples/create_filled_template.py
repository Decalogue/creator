"""
创建填写好的示例 Word 文档用于测试
"""

import os
import sys
from pathlib import Path

***REMOVED*** 添加项目根目录到路径
project_root = Path(__file__).parent.parent.parent.parent
sys.path.insert(0, str(project_root))

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_filled_template(output_path: str):
    """创建填写好的示例文档"""
    
    doc = Document()
    
    ***REMOVED*** 标题
    title = doc.add_heading('短视频创作需求模板', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    ***REMOVED*** 1. 视频基本信息
    doc.add_heading('一、视频基本信息', level=1)
    doc.add_paragraph('视频类型: ecommerce')
    doc.add_paragraph('目标平台: douyin')
    doc.add_paragraph('目标时长: 60')
    doc.add_paragraph()
    
    ***REMOVED*** 2. 当前任务需求
    doc.add_heading('二、当前任务需求', level=1)
    doc.add_paragraph('推广新品手机')
    doc.add_paragraph('突出性价比和拍照功能')
    doc.add_paragraph('目标受众：年轻人')
    doc.add_paragraph('强调真实使用体验')
    doc.add_paragraph()
    
    ***REMOVED*** 3. 修改需求（可选）
    doc.add_heading('三、修改需求（可选）', level=1)
    ***REMOVED*** 留空，这是第一次生成
    doc.add_paragraph()
    
    ***REMOVED*** 4. 通用记忆总结（可选）
    doc.add_heading('四、通用记忆总结（可选）', level=1)
    doc.add_paragraph('喜欢用生活化语言')
    doc.add_paragraph('避免使用"姐妹们"等称呼')
    doc.add_paragraph('偏好真实体验分享')
    doc.add_paragraph('语气要像朋友分享一样自然')
    doc.add_paragraph()
    
    ***REMOVED*** 5. 用户偏好设置
    doc.add_heading('五、用户偏好设置（可选）', level=1)
    doc.add_paragraph('风格偏好: 真诚自然')
    doc.add_paragraph('语气偏好: 像朋友分享')
    doc.add_paragraph('平台偏好: 抖音')
    doc.add_paragraph()
    
    ***REMOVED*** 6. 商品信息（电商题材）
    doc.add_heading('六、商品信息（仅电商题材需要）', level=1)
    doc.add_paragraph('产品名称: 新品手机')
    doc.add_paragraph('核心卖点: 性价比、拍照、快速充电')
    doc.add_paragraph('目标价格: 2000-3000元')
    doc.add_paragraph('特色功能: 夜景模式、超广角镜头')
    doc.add_paragraph()
    
    ***REMOVED*** 7. 镜头素材
    doc.add_heading('七、镜头素材', level=1)
    doc.add_paragraph('镜头1: 产品特写，展示手机外观')
    doc.add_paragraph('镜头2: 使用场景，年轻人拍照')
    doc.add_paragraph('镜头3: 夜景拍摄效果展示')
    doc.add_paragraph('镜头4: 充电速度演示')
    doc.add_paragraph('镜头5: 多场景使用，日常生活')
    doc.add_paragraph()
    
    ***REMOVED*** 保存文档
    doc.save(output_path)
    print(f"填写好的示例文档已创建: {output_path}")

if __name__ == "__main__":
    output_path = "video_script_filled_example.docx"
    create_filled_template(output_path)
    print(f"\n可以运行以下命令生成剧本：")
    print(f"cd /root/data/AI/creator")
    print(f"conda activate seeme")
    print(f"PYTHONPATH=/root/data/AI/creator/src:$PYTHONPATH python src/unimem/examples/generate_video_script.py --doc src/unimem/examples/{output_path} --output generated_script.json")

