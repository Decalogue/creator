"""
创建多个差异化的需求文档用于测试
用于测试不同场景下的记忆积累和模式提取
"""

import os
import sys
from pathlib import Path

***REMOVED*** 添加项目根目录到路径
project_root = Path(__file__).parent.parent.parent.parent
sys.path.insert(0, str(project_root))

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_doc(output_path: str, config: dict):
    """根据配置创建需求文档"""
    
    doc = Document()
    
    ***REMOVED*** 标题
    title = doc.add_heading('短视频创作需求模板', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    ***REMOVED*** 1. 视频基本信息
    doc.add_heading('一、视频基本信息', level=1)
    doc.add_paragraph(f'视频类型: {config.get("video_type", "ecommerce")}')
    doc.add_paragraph(f'目标平台: {config.get("platform", "douyin")}')
    doc.add_paragraph(f'目标时长: {config.get("duration", 60)}')
    doc.add_paragraph()
    
    ***REMOVED*** 2. 当前任务需求
    doc.add_heading('二、当前任务需求', level=1)
    for req in config.get("task_requirements", []):
        doc.add_paragraph(req)
    doc.add_paragraph()
    
    ***REMOVED*** 3. 修改需求（可选）
    doc.add_heading('三、修改需求（可选）', level=1)
    for req in config.get("modification_requirements", []):
        doc.add_paragraph(req)
    doc.add_paragraph()
    
    ***REMOVED*** 4. 通用记忆总结（可选）
    doc.add_heading('四、通用记忆总结（可选）', level=1)
    for mem in config.get("general_memories", []):
        doc.add_paragraph(mem)
    doc.add_paragraph()
    
    ***REMOVED*** 5. 用户偏好设置
    doc.add_heading('五、用户偏好设置（可选）', level=1)
    for pref in config.get("user_preferences", []):
        doc.add_paragraph(pref)
    doc.add_paragraph()
    
    ***REMOVED*** 6. 商品信息（电商题材）
    if config.get("video_type") == "ecommerce":
        doc.add_heading('六、商品信息', level=1)
        for info in config.get("product_info", []):
            doc.add_paragraph(info)
        doc.add_paragraph()
    
    ***REMOVED*** 7. 镜头素材
    doc.add_heading('七、镜头素材', level=1)
    for shot in config.get("shot_materials", []):
        doc.add_paragraph(shot)
    doc.add_paragraph()
    
    ***REMOVED*** 保存文档
    doc.save(output_path)
    print(f"✓ 已创建: {output_path}")

def main():
    """创建多个差异化的需求文档"""
    
    print("="*70)
    print("创建多个差异化的需求文档")
    print("="*70)
    
    ***REMOVED*** 文档1：电商-手机-抖音（已有，作为基准）
    ***REMOVED*** 文档2：电商-美妆-小红书
    config2 = {
        "video_type": "ecommerce",
        "platform": "xiaohongshu",
        "duration": 45,
        "task_requirements": [
            "推广新品口红",
            "突出持久度和显色度",
            "目标受众：年轻女性",
            "强调自然妆容效果"
        ],
        "modification_requirements": [],
        "general_memories": [
            "喜欢真实试色展示",
            "偏好自然光线拍摄",
            "避免过度美颜滤镜"
        ],
        "user_preferences": [
            "风格偏好: 清新自然",
            "语气偏好: 温柔亲切",
            "平台偏好: 小红书"
        ],
        "product_info": [
            "产品名称: 新品口红",
            "核心卖点: 持久、显色、不脱色",
            "目标价格: 100-200元",
            "特色功能: 12小时持久、多种色号"
        ],
        "shot_materials": [
            "镜头1: 产品特写，展示口红外观",
            "镜头2: 试色过程，自然光线",
            "镜头3: 持久度测试，12小时后效果",
            "镜头4: 多色号展示",
            "镜头5: 日常妆容搭配"
        ]
    }
    
    ***REMOVED*** 文档3：知识-教育-抖音
    config3 = {
        "video_type": "education",
        "platform": "douyin",
        "duration": 90,
        "task_requirements": [
            "讲解Python编程基础",
            "面向编程初学者",
            "需要通俗易懂",
            "配合实例演示"
        ],
        "modification_requirements": [],
        "general_memories": [
            "喜欢循序渐进的教学方式",
            "偏好实际案例演示",
            "避免过于理论化"
        ],
        "user_preferences": [
            "风格偏好: 专业但易懂",
            "语气偏好: 耐心细致",
            "平台偏好: 抖音"
        ],
        "product_info": [],
        "shot_materials": [
            "镜头1: 课程标题和讲师介绍",
            "镜头2: 代码演示屏幕录制",
            "镜头3: 实际运行效果展示",
            "镜头4: 重点知识点标注",
            "镜头5: 练习题展示"
        ]
    }
    
    ***REMOVED*** 文档4：娱乐-搞笑-抖音
    config4 = {
        "video_type": "entertainment",
        "platform": "douyin",
        "duration": 30,
        "task_requirements": [
            "制作搞笑短视频",
            "主题：日常生活趣事",
            "需要快速抓住观众注意力",
            "结尾要有反转或笑点"
        ],
        "modification_requirements": [],
        "general_memories": [
            "喜欢真实生活场景",
            "偏好自然表演",
            "避免过度夸张"
        ],
        "user_preferences": [
            "风格偏好: 轻松幽默",
            "语气偏好: 自然随意",
            "平台偏好: 抖音"
        ],
        "product_info": [],
        "shot_materials": [
            "镜头1: 日常场景，制造悬念",
            "镜头2: 意外情况发生",
            "镜头3: 反应镜头",
            "镜头4: 反转或笑点",
            "镜头5: 结尾总结"
        ]
    }
    
    ***REMOVED*** 文档5：电商-服装-淘宝
    config5 = {
        "video_type": "ecommerce",
        "platform": "taobao",
        "duration": 60,
        "task_requirements": [
            "推广新款连衣裙",
            "突出面料和版型",
            "目标受众：25-35岁女性",
            "强调多场景穿搭"
        ],
        "modification_requirements": [],
        "general_memories": [
            "喜欢详细展示细节",
            "偏好真实上身效果",
            "避免过度修图"
        ],
        "user_preferences": [
            "风格偏好: 优雅知性",
            "语气偏好: 专业推荐",
            "平台偏好: 淘宝"
        ],
        "product_info": [
            "产品名称: 新款连衣裙",
            "核心卖点: 优质面料、修身版型、多场景适用",
            "目标价格: 300-500元",
            "特色功能: 易打理、不易起皱、四季可穿"
        ],
        "shot_materials": [
            "镜头1: 产品整体展示",
            "镜头2: 面料细节特写",
            "镜头3: 上身效果展示",
            "镜头4: 多场景穿搭演示",
            "镜头5: 细节做工展示"
        ]
    }
    
    configs = [
        ("video_script_test1_ecommerce_phone_douyin.docx", config2),  ***REMOVED*** 使用不同的配置
        ("video_script_test2_ecommerce_lipstick_xiaohongshu.docx", config2),
        ("video_script_test3_education_python_douyin.docx", config3),
        ("video_script_test4_entertainment_funny_douyin.docx", config4),
        ("video_script_test5_ecommerce_dress_taobao.docx", config5),
    ]
    
    for filename, config in configs:
        create_doc(filename, config)
    
    print(f"\n✓ 已创建 {len(configs)} 个差异化的需求文档")
    print("\n文档列表：")
    for filename, _ in configs:
        print(f"  - {filename}")
    
    print("\n" + "="*70)
    print("可以开始运行多个测试流程了")
    print("="*70)

if __name__ == "__main__":
    main()
