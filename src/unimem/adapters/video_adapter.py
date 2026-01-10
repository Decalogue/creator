"""
çŸ­è§†é¢‘æ–‡æ¡ˆé€‚é…å™¨

ä¸“é—¨ç”¨äºè‡ªåª’ä½“å’Œç”µå•†çŸ­è§†é¢‘æ–‡æ¡ˆåˆ›ä½œçš„é€‚é…å™¨ï¼ŒåŸºäº atom_link_adapter çš„åŸºç¡€æ¶æ„ï¼Œ
é’ˆå¯¹çŸ­è§†é¢‘æ··å‰ªåˆ›ä½œçš„ç‰¹æ®Šéœ€æ±‚è¿›è¡Œå®šåˆ¶åŒ–å®ç°ã€‚

çŸ­è§†é¢‘åˆ›ä½œç‰¹ç‚¹ï¼š
- å‰ªè¾‘æµç¨‹ï¼šç´ æå¯¼å…¥ -> æ–‡æ¡ˆæ’°å†™ -> é•œå¤´åŒ¹é… -> èŠ‚å¥æ§åˆ¶ -> æœ€ç»ˆæˆç‰‡
- ä¸ªæ€§åŒ–å®šåˆ¶ï¼šåŸºäºç”¨æˆ·è®°å¿†å’Œåå¥½
- å•†å“è¥é”€ï¼šç»“åˆå•†å“ä¿¡æ¯è¿›è¡Œæ¨å¹¿
- ç´ æåˆ©ç”¨ï¼šæœ€å¤§åŒ–åˆ©ç”¨æ‹æ‘„é•œå¤´ç´ æ

ä¸»è¦åŠŸèƒ½ï¼š
1. Word æ–‡æ¡£è§£æï¼šè§£æç”¨æˆ·è®°å¿†ã€å•†å“ä¿¡æ¯ã€é•œå¤´ç´ æ
2. ä¸ªæ€§åŒ–æ–‡æ¡ˆç”Ÿæˆï¼šåŸºäºç”¨æˆ·åå¥½ç”Ÿæˆæ–‡æ¡ˆ
3. é•œå¤´ç´ æåŒ¹é…ï¼šæ ¹æ®æ–‡æ¡ˆåŒ¹é…æœ€ä½³é•œå¤´
4. å‰ªè¾‘è„šæœ¬ç”Ÿæˆï¼šç”Ÿæˆå®Œæ•´çš„å‰ªè¾‘è„šæœ¬
5. èŠ‚å¥æ§åˆ¶ï¼šæ§åˆ¶è§†é¢‘èŠ‚å¥å’Œæ—¶é•¿

å·¥ä¸šçº§ç‰¹æ€§ï¼š
- å‚æ•°éªŒè¯å’Œè¾“å…¥æ£€æŸ¥
- ç»Ÿä¸€å¼‚å¸¸å¤„ç†
- Word æ–‡æ¡£è§£ææ”¯æŒ
- é”™è¯¯å¤„ç†å’Œé™çº§
"""

import json
import logging
from typing import Dict, Any, Optional, List
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

from .atom_link_adapter import AtomLinkAdapter
from .base import AdapterError, AdapterNotAvailableError
from ..memory_types import Memory, Experience, Context, MemoryType
from ..chat import ark_deepseek_v3_2

logger = logging.getLogger(__name__)


class VideoAdapter(AtomLinkAdapter):
    """
    çŸ­è§†é¢‘æ–‡æ¡ˆé€‚é…å™¨
    
    ç»§æ‰¿ AtomLinkAdapter çš„åŸºç¡€åŠŸèƒ½ï¼Œä½†é’ˆå¯¹çŸ­è§†é¢‘æ··å‰ªåˆ›ä½œè¿›è¡Œå®šåˆ¶åŒ–ã€‚
    
    æ ¸å¿ƒåŠŸèƒ½ï¼š
    1. Word æ–‡æ¡£è§£æï¼šè§£æç”¨æˆ·è®°å¿†ã€å•†å“ä¿¡æ¯ã€é•œå¤´ç´ æ
    2. ä¸ªæ€§åŒ–æ–‡æ¡ˆç”Ÿæˆï¼šåŸºäºç”¨æˆ·åå¥½ç”Ÿæˆæ–‡æ¡ˆ
    3. é•œå¤´ç´ æåŒ¹é…ï¼šæ ¹æ®æ–‡æ¡ˆåŒ¹é…æœ€ä½³é•œå¤´
    4. å‰ªè¾‘è„šæœ¬ç”Ÿæˆï¼šç”Ÿæˆå®Œæ•´çš„å‰ªè¾‘è„šæœ¬
    
    UniMem ä¼˜åŠ¿åˆ©ç”¨ï¼š
    1. è¯­ä¹‰æ£€ç´¢ï¼šä» UniMem ä¸­æ£€ç´¢ç›¸å…³çš„å†å²åˆ›ä½œå’Œç”¨æˆ·åå¥½
    2. è®°å¿†å­˜å‚¨ï¼šå°†ç”Ÿæˆçš„è„šæœ¬å­˜å‚¨åˆ° UniMemï¼Œå»ºç«‹å…³è”ç½‘ç»œ
    3. è·¨ä»»åŠ¡å­¦ä¹ ï¼šåˆ©ç”¨é€šç”¨è®°å¿†å’Œä»»åŠ¡è®°å¿†ï¼Œå®ç°é£æ ¼ä¸€è‡´æ€§
    4. è®°å¿†æ¼”åŒ–ï¼šé€šè¿‡ REFLECT æ“ä½œä¼˜åŒ–å’Œæ€»ç»“åˆ›ä½œç»éªŒ
    """
    
    def __init__(self, config: Optional[Dict[str, Any]] = None, unimem_instance=None):
        """
        åˆå§‹åŒ–çŸ­è§†é¢‘æ–‡æ¡ˆé€‚é…å™¨
        
        Args:
            config: é€‚é…å™¨é…ç½®
            unimem_instance: UniMem å®ä¾‹ï¼ˆå¯é€‰ï¼‰ï¼Œç”¨äºæ·±åº¦é›†æˆ UniMem åŠŸèƒ½
        """
        super().__init__(config)
        self.unimem = unimem_instance
        if self.unimem:
            logger.info("Video adapter initialized with UniMem integration")
        else:
            logger.info("Video adapter initialized (standalone mode)")
    
    def _do_initialize(self) -> None:
        """
        åˆå§‹åŒ–çŸ­è§†é¢‘æ–‡æ¡ˆé€‚é…å™¨
        
        è°ƒç”¨çˆ¶ç±»åˆå§‹åŒ–ä»¥å¤ç”¨åŸºç¡€æ¶æ„ï¼ˆå‘é‡å­˜å‚¨ã€åµŒå…¥æ¨¡å‹ç­‰ï¼‰ï¼Œ
        ç„¶åè¿›è¡ŒçŸ­è§†é¢‘åˆ›ä½œç‰¹å®šçš„åˆå§‹åŒ–ã€‚
        """
        super()._do_initialize()
        logger.info("Video adapter initialized (specialized for short video script writing)")
    
    def enrich_memories_from_unimem(
        self,
        task_memories: List[str],
        video_type: str,
        top_k: int = 10
    ) -> Dict[str, List[str]]:
        """
        ä» UniMem ä¸­ä¸°å¯Œè®°å¿†ä¿¡æ¯
        
        åˆ©ç”¨ UniMem çš„è¯­ä¹‰æ£€ç´¢èƒ½åŠ›ï¼Œä»å†å²è®°å¿†ä¸­æ£€ç´¢ï¼š
        1. ç›¸å…³çš„å†å²åˆ›ä½œï¼ˆç›¸ä¼¼é¢˜æã€é£æ ¼ã€ä¸»é¢˜ï¼‰
        2. ç”¨æˆ·çš„åˆ›ä½œåå¥½å’Œä¹ æƒ¯
        3. æˆåŠŸçš„åˆ›ä½œæ¨¡å¼å’Œç»éªŒ
        4. ç›¸å…³çš„å®ä½“å’Œå…³ç³»
        
        Args:
            task_memories: å½“å‰ä»»åŠ¡è®°å¿†åˆ—è¡¨
            video_type: è§†é¢‘ç±»å‹
            top_k: æ¯ä¸ªæ£€ç´¢ç»´åº¦è¿”å›çš„ç»“æœæ•°é‡
            
        Returns:
            ä¸°å¯Œçš„è®°å¿†å­—å…¸ï¼š
            - historical_scripts: å†å²ç›¸å…³è„šæœ¬
            - successful_patterns: æˆåŠŸçš„åˆ›ä½œæ¨¡å¼
            - related_entities: ç›¸å…³çš„å®ä½“å’Œå…³ç³»
            - user_style_preferences: ç”¨æˆ·çš„é£æ ¼åå¥½
        """
        if not self.unimem:
            logger.debug("UniMem not available, skipping memory enrichment")
            return {
                "historical_scripts": [],
                "successful_patterns": [],
                "related_entities": [],
                "user_style_preferences": []
            }
        
        enriched = {
            "historical_scripts": [],
            "successful_patterns": [],
            "related_entities": [],
            "user_style_preferences": []
        }
        
        try:
            ***REMOVED*** 1. ä»ä»»åŠ¡è®°å¿†æ„å»ºæŸ¥è¯¢ï¼Œæ£€ç´¢ç›¸å…³çš„å†å²åˆ›ä½œ
            task_query = " ".join(task_memories[:3])  ***REMOVED*** ä½¿ç”¨å‰3ä¸ªä»»åŠ¡è®°å¿†ä½œä¸ºæŸ¥è¯¢
            if task_query:
                ***REMOVED*** ä½¿ç”¨ UniMem çš„è¯­ä¹‰æ£€ç´¢
                ***REMOVED*** æ£€æŸ¥æ˜¯å¦æ˜¯ HTTP æœåŠ¡å®¢æˆ·ç«¯ï¼ˆä¸éœ€è¦ context å‚æ•°ï¼‰
                if hasattr(self.unimem, 'base_url'):
                    recall_results = self.unimem.recall(
                        query=f"{video_type} {task_query}",
                        top_k=top_k
                    )
                else:
                    recall_results = self.unimem.recall(
                        query=f"{video_type} {task_query}",
                        context=Context(),
                        top_k=top_k
                    )
                
                for result in recall_results:
                    if result.memory:
                        ***REMOVED*** ç­›é€‰è§†é¢‘ç›¸å…³çš„è®°å¿†
                        content = result.memory.content
                        if any(keyword in content.lower() for keyword in 
                               ["è§†é¢‘", "è„šæœ¬", "æ–‡æ¡ˆ", "é•œå¤´", "å‰ªè¾‘", "video", "script"]):
                            enriched["historical_scripts"].append({
                                "content": content[:200],  ***REMOVED*** æˆªå–å‰200å­—ç¬¦
                                "score": result.score,
                                "metadata": result.memory.metadata
                            })
            
            ***REMOVED*** 2. æ£€ç´¢æˆåŠŸçš„åˆ›ä½œæ¨¡å¼ï¼ˆé€šè¿‡æ ‡ç­¾æˆ–å…ƒæ•°æ®ï¼‰
            if video_type:
                if hasattr(self.unimem, 'base_url'):
                    pattern_results = self.unimem.recall(
                        query=f"æˆåŠŸ {video_type} æ¨¡å¼ ç»éªŒ",
                        top_k=top_k
                    )
                else:
                    pattern_results = self.unimem.recall(
                        query=f"æˆåŠŸ {video_type} æ¨¡å¼ ç»éªŒ",
                        context=Context(),
                        top_k=top_k
                    )
                for result in pattern_results[:5]:  ***REMOVED*** åªå–å‰5ä¸ª
                    if result.memory:
                        enriched["successful_patterns"].append({
                            "pattern": result.memory.content[:150],
                            "score": result.score
                        })
            
            ***REMOVED*** 3. æ£€ç´¢ç”¨æˆ·é£æ ¼åå¥½ï¼ˆé€šè¿‡é€šç”¨è®°å¿†ï¼‰
            if hasattr(self.unimem, 'base_url'):
                style_results = self.unimem.recall(
                    query="ç”¨æˆ· é£æ ¼ åå¥½ è¯­æ°” è¡¨è¾¾æ–¹å¼",
                    top_k=top_k
                )
            else:
                style_results = self.unimem.recall(
                    query="ç”¨æˆ· é£æ ¼ åå¥½ è¯­æ°” è¡¨è¾¾æ–¹å¼",
                    context=Context(),
                    top_k=top_k
                )
            for result in style_results[:5]:
                if result.memory:
                    enriched["user_style_preferences"].append({
                        "preference": result.memory.content[:150],
                        "score": result.score
                    })
            
            logger.info(f"Enriched memories from UniMem: {len(enriched['historical_scripts'])} scripts, "
                       f"{len(enriched['successful_patterns'])} patterns, "
                       f"{len(enriched['user_style_preferences'])} preferences")
        except Exception as e:
            logger.warning(f"Error enriching memories from UniMem: {e}", exc_info=True)
        
        return enriched
    
    def store_script_to_unimem(
        self,
        script_data: Dict[str, Any],
        task_memories: List[str],
        video_type: str,
        platform: str
    ) -> Optional[str]:
        """
        å°†ç”Ÿæˆçš„è„šæœ¬å­˜å‚¨åˆ° UniMem
        
        ä½¿ç”¨ UniMem çš„ RETAIN æ“ä½œï¼Œå°†è„šæœ¬å†…å®¹å­˜å‚¨ä¸ºè®°å¿†ï¼Œ
        å¹¶å»ºç«‹ä¸å…¶ä»–è®°å¿†çš„å…³è”ï¼Œä¾¿äºåç»­æ£€ç´¢å’Œå­¦ä¹ ã€‚
        
        Args:
            script_data: ç”Ÿæˆçš„è„šæœ¬æ•°æ®
            task_memories: ä»»åŠ¡è®°å¿†åˆ—è¡¨
            video_type: è§†é¢‘ç±»å‹
            platform: å¹³å°ç±»å‹
            
        Returns:
            å­˜å‚¨çš„è®°å¿† IDï¼Œå¦‚æœå­˜å‚¨å¤±è´¥è¿”å› None
        """
        if not self.unimem:
            logger.debug("UniMem not available, skipping script storage")
            return None
        
        try:
            ***REMOVED*** æ„å»ºè®°å¿†å†…å®¹
            script_summary = script_data.get("summary", {})
            script_content = script_data.get("script", "")
            
            ***REMOVED*** åˆ›å»º Experience å¯¹è±¡
            experience_content = f"""
è§†é¢‘ç±»å‹ï¼š{video_type}
å¹³å°ï¼š{platform}
è„šæœ¬æ‘˜è¦ï¼š{json.dumps(script_summary, ensure_ascii=False)}
è„šæœ¬é¢„è§ˆï¼š{script_content[:500]}
ä»»åŠ¡éœ€æ±‚ï¼š{' | '.join(task_memories[:3])}
            """.strip()
            
            experience = Experience(
                content=experience_content,
                timestamp=datetime.now()
            )
            
            ***REMOVED*** åˆ›å»º Context å¯¹è±¡
            ***REMOVED*** æ„å»ºå†³ç­–ä¸Šä¸‹æ–‡ï¼ˆContext Graphå¢å¼ºï¼‰
            context = Context(
                metadata={
                    "source": "video_script",  ***REMOVED*** æ˜ç¡®æ ‡è¯†æ¥æº
                    "task_description": f"ç”Ÿæˆ{video_type}ç±»å‹çŸ­è§†é¢‘è„šæœ¬",
                    "video_type": video_type,
                    "platform": platform,
                    "segments_count": len(script_data.get("segments", [])),
                    "duration": script_data.get("editing_script", {}).get("total_duration", 0),
                    ***REMOVED*** å†³ç­–ç—•è¿¹ï¼ˆContext Graphå¢å¼ºï¼‰
                    "inputs": task_memories[:5] if task_memories else [],  ***REMOVED*** ä½¿ç”¨çš„ä»»åŠ¡è®°å¿†
                    "rules": [
                        f"{platform}å¹³å°è§„åˆ™",
                        f"{video_type}ç±»å‹è„šæœ¬è§„èŒƒ",
                        "3ç§’åŸåˆ™ï¼ˆå‰3ç§’æŠ“ä½æ³¨æ„åŠ›ï¼‰",
                        "è½¬åŒ–ç‡ä¼˜åŒ–è¦æ±‚"
                    ],
                    "exceptions": [],  ***REMOVED*** å¯ä»¥åœ¨è¿™é‡Œæ·»åŠ å¼‚å¸¸æƒ…å†µ
                    "approvals": [],  ***REMOVED*** å¯ä»¥åœ¨è¿™é‡Œæ·»åŠ å®¡æ‰¹æµç¨‹
                    "reasoning": f"åŸºäº{len(task_memories) if task_memories else 0}æ¡ä»»åŠ¡è®°å¿†ï¼Œç”Ÿæˆ{video_type}ç±»å‹è„šæœ¬ï¼Œé€‚é…{platform}å¹³å°ç‰¹ç‚¹",
                    ***REMOVED*** ç¡®ä¿decision_traceå­—æ®µè¢«æ­£ç¡®è®¾ç½®
                    "decision_trace": {
                        "inputs": task_memories[:5] if task_memories else [],
                        "rules_applied": [
                            f"{platform}å¹³å°è§„åˆ™",
                            f"{video_type}ç±»å‹è„šæœ¬è§„èŒƒ",
                            "3ç§’åŸåˆ™ï¼ˆå‰3ç§’æŠ“ä½æ³¨æ„åŠ›ï¼‰",
                            "è½¬åŒ–ç‡ä¼˜åŒ–è¦æ±‚"
                        ],
                        "exceptions": [],
                        "approvals": [],
                        "timestamp": datetime.now().isoformat(),
                    }
                }
            )
            
            ***REMOVED*** ä½¿ç”¨ UniMem çš„ RETAIN æ“ä½œå­˜å‚¨
            memory = self.unimem.retain(experience, context)
            
            logger.info(f"Stored script to UniMem: memory_id={memory.id}")
            return memory.id
        except Exception as e:
            logger.warning(f"Error storing script to UniMem: {e}", exc_info=True)
            return None
    
    def create_word_template(
        self, 
        output_path: str = "video_script_template.docx",
        auto_fill_from_memory: bool = True
    ) -> str:
        """
        åˆ›å»º Word æ¨¡æ¿æ–‡æ¡£ï¼Œç”¨æˆ·å¯ä»¥ä¸‹è½½å¡«å†™åä¸Šä¼ 
        
        æ¨¡æ¿åŒ…å«æ‰€æœ‰å¿…è¦çš„å­—æ®µï¼Œç”¨æˆ·åªéœ€å¡«å†™ç›¸åº”å†…å®¹å³å¯ã€‚
        
        å¢å¼ºåŠŸèƒ½ï¼š
        1. æ™ºèƒ½æç¤ºï¼šæ ¹æ®è§†é¢‘ç±»å‹æä¾›ä¸åŒçš„å¡«å†™å»ºè®®
        2. è‡ªåŠ¨å¡«å……ï¼šä»UniMemå†å²è®°å¿†ä¸­è‡ªåŠ¨å¡«å……é€šç”¨åå¥½å’Œé•œå¤´ç´ æå»ºè®®
        3. æ ¼å¼å®¹é”™ï¼šæ”¯æŒæ›´å¤šçµæ´»çš„å¡«å†™æ ¼å¼
        
        Args:
            output_path: æ¨¡æ¿æ–‡ä»¶ä¿å­˜è·¯å¾„ï¼ˆé»˜è®¤ä¸º "video_script_template.docx"ï¼‰
            auto_fill_from_memory: æ˜¯å¦ä»å†å²è®°å¿†ä¸­è‡ªåŠ¨å¡«å……éƒ¨åˆ†å†…å®¹ï¼ˆé»˜è®¤Trueï¼‰
            
        Returns:
            åˆ›å»ºçš„æ¨¡æ¿æ–‡ä»¶è·¯å¾„
            
        Raises:
            AdapterNotAvailableError: å¦‚æœ python-docx ä¸å¯ç”¨
            AdapterError: å¦‚æœè·¯å¾„æ— æ•ˆ
        """
        if not DOCX_AVAILABLE:
            raise AdapterNotAvailableError(
                "python-docx library not available. Install with: pip install python-docx",
                adapter_name="VideoAdapter"
            )
        
        if not output_path or not isinstance(output_path, str):
            raise AdapterError("output_path must be a non-empty string", adapter_name="VideoAdapter")
        
        try:
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            ***REMOVED*** æ ‡é¢˜
            title = doc.add_heading('çŸ­è§†é¢‘åˆ›ä½œéœ€æ±‚æ¨¡æ¿', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            ***REMOVED*** è¯´æ˜
            doc.add_paragraph('è¯·æ ¹æ®ä»¥ä¸‹è¯´æ˜å¡«å†™ç›¸åº”å†…å®¹ï¼Œæ‰€æœ‰å†…å®¹éƒ½æ˜¯å¯é€‰çš„ï¼Œæ ¹æ®éœ€è¦å¡«å†™å³å¯ã€‚')
            
            ***REMOVED*** æ™ºèƒ½æç¤ºè¯´æ˜
            if auto_fill_from_memory and self.unimem:
                doc.add_paragraph('ğŸ’¡ æç¤ºï¼šç³»ç»Ÿå·²æ ¹æ®æ‚¨çš„å†å²åå¥½è‡ªåŠ¨å¡«å……äº†éƒ¨åˆ†å»ºè®®å†…å®¹ï¼Œæ‚¨å¯ä»¥ç›´æ¥ä½¿ç”¨æˆ–ä¿®æ”¹ã€‚')
            else:
                doc.add_paragraph('ğŸ’¡ æç¤ºï¼šå»ºè®®å¡«å†™å°½å¯èƒ½è¯¦ç»†çš„ä¿¡æ¯ï¼Œä»¥ä¾¿ç”Ÿæˆæ›´ç¬¦åˆæ‚¨éœ€æ±‚çš„è§†é¢‘å‰§æœ¬ã€‚')
            doc.add_paragraph()
            
            ***REMOVED*** 1. è§†é¢‘åŸºæœ¬ä¿¡æ¯
            doc.add_heading('ä¸€ã€è§†é¢‘åŸºæœ¬ä¿¡æ¯', level=1)
            doc.add_paragraph('è§†é¢‘ç±»å‹ï¼ˆå¿…å¡«ï¼Œä»ä»¥ä¸‹é€‰é¡¹é€‰æ‹©å…¶ä¸€ï¼‰ï¼š')
            doc.add_paragraph('  - ecommerceï¼ˆç”µå•†æ¨å¹¿ï¼‰', style='List Bullet')
            doc.add_paragraph('  - ip_buildingï¼ˆä¸ªäººIPæ‰“é€ ï¼‰', style='List Bullet')
            doc.add_paragraph('  - knowledgeï¼ˆçŸ¥è¯†åˆ†äº«ï¼‰', style='List Bullet')
            doc.add_paragraph('  - vlogï¼ˆç”Ÿæ´»Vlogï¼‰', style='List Bullet')
            doc.add_paragraph('  - mediaï¼ˆè‡ªåª’ä½“å†…å®¹ï¼‰', style='List Bullet')
            doc.add_paragraph('è§†é¢‘ç±»å‹: [è¯·å¡«å†™ï¼Œä¾‹å¦‚ï¼šecommerce]')
            doc.add_paragraph()
            
            doc.add_paragraph('ç›®æ ‡å¹³å°ï¼ˆå¿…å¡«ï¼Œä»ä»¥ä¸‹é€‰é¡¹é€‰æ‹©å…¶ä¸€ï¼‰ï¼š')
            doc.add_paragraph('  - douyinï¼ˆæŠ–éŸ³ï¼‰', style='List Bullet')
            doc.add_paragraph('  - xiaohongshuï¼ˆå°çº¢ä¹¦ï¼‰', style='List Bullet')
            doc.add_paragraph('  - tiktokï¼ˆTikTokå›½é™…ï¼‰', style='List Bullet')
            doc.add_paragraph('  - youtubeï¼ˆYouTubeï¼‰', style='List Bullet')
            doc.add_paragraph('ç›®æ ‡å¹³å°: [è¯·å¡«å†™ï¼Œä¾‹å¦‚ï¼šdouyin]')
            doc.add_paragraph()
            
            doc.add_paragraph('ç›®æ ‡æ—¶é•¿ï¼ˆç§’ï¼Œå¯é€‰ï¼Œé»˜è®¤60ç§’ï¼‰: [è¯·å¡«å†™æ•°å­—ï¼Œä¾‹å¦‚ï¼š60]')
            doc.add_paragraph()
            
            ***REMOVED*** 2. å½“å‰ä»»åŠ¡éœ€æ±‚
            doc.add_heading('äºŒã€å½“å‰ä»»åŠ¡éœ€æ±‚', level=1)
            doc.add_paragraph('è¯·è¯¦ç»†æè¿°æœ¬æ¬¡è§†é¢‘åˆ›ä½œçš„å…·ä½“éœ€æ±‚ï¼Œæ¯è¡Œä¸€æ¡ï¼š')
            doc.add_paragraph('[ä¾‹å¦‚ï¼šæ¨å¹¿æ–°å“æ‰‹æœº]')
            doc.add_paragraph('[ä¾‹å¦‚ï¼šçªå‡ºæ€§ä»·æ¯”å’Œæ‹ç…§åŠŸèƒ½]')
            doc.add_paragraph('[ä¾‹å¦‚ï¼šç›®æ ‡å—ä¼—ï¼šå¹´è½»äºº]')
            doc.add_paragraph()
            
            ***REMOVED*** 3. ä¿®æ”¹éœ€æ±‚ï¼ˆå¯é€‰ï¼‰
            doc.add_heading('ä¸‰ã€ä¿®æ”¹éœ€æ±‚ï¼ˆå¯é€‰ï¼‰', level=1)
            doc.add_paragraph('å¦‚æœåœ¨å·²æœ‰è„šæœ¬åŸºç¡€ä¸Šè¿›è¡Œä¿®æ”¹ï¼Œè¯·å¡«å†™ä¿®æ”¹è¦æ±‚ï¼š')
            doc.add_paragraph('[ä¾‹å¦‚ï¼šå¢åŠ æƒ…æ„Ÿå…±é¸£]')
            doc.add_paragraph('[ä¾‹å¦‚ï¼šè°ƒæ•´è¯­æ°”æ›´è½»æ¾]')
            doc.add_paragraph()
            
            ***REMOVED*** 4. é€šç”¨è®°å¿†æ€»ç»“ï¼ˆå¯é€‰ï¼‰
            doc.add_heading('å››ã€é€šç”¨è®°å¿†æ€»ç»“ï¼ˆå¯é€‰ï¼‰', level=1)
            doc.add_paragraph('è·¨ä»»åŠ¡çš„ç”¨æˆ·åå¥½å’Œé€šç”¨é£æ ¼åå¥½ï¼Œè¿™äº›ä¼šåº”ç”¨åˆ°æ‰€æœ‰è§†é¢‘åˆ›ä½œï¼š')
            
            ***REMOVED*** å°è¯•ä»å†å²è®°å¿†ä¸­è‡ªåŠ¨å¡«å……é€šç”¨åå¥½
            auto_filled_general_memories = []
            if auto_fill_from_memory and self.unimem:
                try:
                    auto_filled_general_memories = self._get_auto_fill_general_memories()
                    if auto_filled_general_memories:
                        doc.add_paragraph('ã€ç³»ç»Ÿå»ºè®®ï¼ˆåŸºäºå†å²è®°å½•ï¼‰ã€‘', style='Strong')
                        for mem in auto_filled_general_memories[:5]:
                            doc.add_paragraph(f'âœ“ {mem}', style='List Bullet')
                        doc.add_paragraph('ã€æ‚¨ä¹Ÿå¯ä»¥å¡«å†™æˆ–ä¿®æ”¹ã€‘', style='Strong')
                except Exception as e:
                    logger.debug(f"Failed to auto-fill general memories: {e}")
            
            if not auto_filled_general_memories:
                doc.add_paragraph('[ä¾‹å¦‚ï¼šå–œæ¬¢ç”¨ç”Ÿæ´»åŒ–è¯­è¨€]')
                doc.add_paragraph('[ä¾‹å¦‚ï¼šé¿å…ä½¿ç”¨"å§å¦¹ä»¬"ç­‰ç§°å‘¼]')
                doc.add_paragraph('[ä¾‹å¦‚ï¼šåå¥½çœŸå®ä½“éªŒåˆ†äº«]')
            doc.add_paragraph()
            
            ***REMOVED*** 5. ç”¨æˆ·åå¥½è®¾ç½®
            doc.add_heading('äº”ã€ç”¨æˆ·åå¥½è®¾ç½®ï¼ˆå¯é€‰ï¼‰', level=1)
            doc.add_paragraph('è¯·ä½¿ç”¨"é”®: å€¼"çš„æ ¼å¼å¡«å†™ï¼Œä¾‹å¦‚ï¼š')
            
            ***REMOVED*** å°è¯•ä»å†å²è®°å¿†ä¸­è‡ªåŠ¨å¡«å……ç”¨æˆ·åå¥½
            auto_filled_preferences = {}
            if auto_fill_from_memory and self.unimem:
                try:
                    auto_filled_preferences = self._get_auto_fill_preferences()
                    if auto_filled_preferences:
                        doc.add_paragraph('ã€ç³»ç»Ÿå»ºè®®ï¼ˆåŸºäºå†å²è®°å½•ï¼‰ã€‘', style='Strong')
                        for key, value in list(auto_filled_preferences.items())[:5]:
                            doc.add_paragraph(f'âœ“ {key}: {value}', style='List Bullet')
                        doc.add_paragraph('ã€æ‚¨ä¹Ÿå¯ä»¥å¡«å†™æˆ–ä¿®æ”¹ã€‘', style='Strong')
                except Exception as e:
                    logger.debug(f"Failed to auto-fill preferences: {e}")
            
            if not auto_filled_preferences:
                doc.add_paragraph('é£æ ¼åå¥½: çœŸè¯šè‡ªç„¶')
                doc.add_paragraph('å¹³å°åå¥½: æŠ–éŸ³')
                doc.add_paragraph('è¯­æ°”åå¥½: åƒæœ‹å‹åˆ†äº«')
            doc.add_paragraph('[è¯·åœ¨æ­¤å¡«å†™æ‚¨çš„åå¥½è®¾ç½®]')
            doc.add_paragraph()
            
            ***REMOVED*** 6. å•†å“ä¿¡æ¯ï¼ˆä»…ç”µå•†é¢˜æéœ€è¦ï¼‰
            doc.add_heading('å…­ã€å•†å“ä¿¡æ¯ï¼ˆä»…ç”µå•†é¢˜æéœ€è¦ï¼‰', level=1)
            doc.add_paragraph('è¯·ä½¿ç”¨"é”®: å€¼"çš„æ ¼å¼å¡«å†™ï¼Œä¾‹å¦‚ï¼š')
            doc.add_paragraph('äº§å“åç§°: æ–°å“æ‰‹æœº')
            doc.add_paragraph('æ ¸å¿ƒå–ç‚¹: æ€§ä»·æ¯”ã€æ‹ç…§')
            doc.add_paragraph('ç›®æ ‡ä»·æ ¼: 2000-3000å…ƒ')
            doc.add_paragraph('[è¯·åœ¨æ­¤å¡«å†™å•†å“ä¿¡æ¯]')
            doc.add_paragraph()
            
            ***REMOVED*** 7. é•œå¤´ç´ æ
            doc.add_heading('ä¸ƒã€é•œå¤´ç´ æ', level=1)
            doc.add_paragraph('è¯·æè¿°å¯ç”¨çš„é•œå¤´ç´ æï¼Œæ¯è¡Œä¸€ä¸ªé•œå¤´ï¼Œä½¿ç”¨"é”®: å€¼"æ ¼å¼æˆ–ç›´æ¥æè¿°ï¼š')
            doc.add_paragraph('ğŸ’¡ æç¤ºï¼šé•œå¤´ç´ æè¶Šè¯¦ç»†ï¼Œç”Ÿæˆçš„å‰§æœ¬è¶Šç²¾å‡†ã€‚å»ºè®®åŒ…æ‹¬ï¼šäº§å“å±•ç¤ºã€ä½¿ç”¨åœºæ™¯ã€ç»†èŠ‚ç‰¹å†™ã€å¯¹æ¯”æ•ˆæœç­‰ã€‚')
            
            ***REMOVED*** å°è¯•ä»å†å²è®°å¿†ä¸­è‡ªåŠ¨å¡«å……é•œå¤´ç´ æå»ºè®®
            auto_filled_shots = []
            if auto_fill_from_memory and self.unimem:
                try:
                    auto_filled_shots = self._get_auto_fill_shot_suggestions()
                    if auto_filled_shots:
                        doc.add_paragraph('ã€ç³»ç»Ÿå»ºè®®ï¼ˆåŸºäºå†å²è®°å½•ï¼‰ã€‘', style='Strong')
                        for i, shot in enumerate(auto_filled_shots[:8], 1):
                            doc.add_paragraph(f'âœ“ é•œå¤´{i}: {shot}', style='List Bullet')
                        doc.add_paragraph('ã€æ‚¨ä¹Ÿå¯ä»¥å¡«å†™æˆ–ä¿®æ”¹ã€‘', style='Strong')
                except Exception as e:
                    logger.debug(f"Failed to auto-fill shot suggestions: {e}")
            
            if not auto_filled_shots:
                doc.add_paragraph('é•œå¤´1: äº§å“ç‰¹å†™ï¼Œå±•ç¤ºæ‰‹æœºå¤–è§‚')
                doc.add_paragraph('é•œå¤´2: ä½¿ç”¨åœºæ™¯ï¼Œå¹´è½»äººæ‹ç…§')
                doc.add_paragraph('[è¯·åœ¨æ­¤å¡«å†™é•œå¤´ç´ æ]')
            doc.add_paragraph()
            
            ***REMOVED*** ä¿å­˜æ–‡æ¡£
            doc.save(output_path)
            logger.info(f"Word template created: {output_path} (auto_fill: {auto_fill_from_memory})")
            return output_path
        except Exception as e:
            raise AdapterError(
                f"Failed to create Word template: {e}",
                adapter_name="VideoAdapter",
                cause=e
            ) from e
    
    def _get_auto_fill_general_memories(self) -> List[str]:
        """
        ä»UniMemå†å²è®°å¿†ä¸­è·å–é€šç”¨è®°å¿†å»ºè®®
        
        Returns:
            é€šç”¨è®°å¿†åˆ—è¡¨
        """
        if not self.unimem:
            return []
        
        try:
            ***REMOVED*** æ£€ç´¢é€šç”¨è®°å¿†å’Œé£æ ¼åå¥½
            results = self.unimem.recall(
                query="ç”¨æˆ·é£æ ¼åå¥½ é€šç”¨è®°å¿† åˆ›ä½œåå¥½",
                memory_type=MemoryType.OPINION,  ***REMOVED*** åå¥½é€šå¸¸æ˜¯OPINIONç±»å‹
                top_k=5
            )
            
            memories = []
            for result in results:
                if result.memory and result.memory.content:
                    content = result.memory.content
                    ***REMOVED*** æå–å…³é”®åå¥½ä¿¡æ¯ï¼ˆè¿‡æ»¤æ‰å¤ªé•¿çš„å†…å®¹ï¼‰
                    if len(content) < 200:
                        memories.append(content)
            
            return memories[:5]
        except Exception as e:
            logger.debug(f"Failed to get auto-fill general memories: {e}")
            return []
    
    def _get_auto_fill_preferences(self) -> Dict[str, str]:
        """
        ä»UniMemå†å²è®°å¿†ä¸­è·å–ç”¨æˆ·åå¥½å»ºè®®
        
        Returns:
            åå¥½å­—å…¸
        """
        if not self.unimem:
            return {}
        
        try:
            ***REMOVED*** æ£€ç´¢ç”¨æˆ·åå¥½è®°å¿†
            results = self.unimem.recall(
                query="ç”¨æˆ·åå¥½è®¾ç½® é£æ ¼åå¥½ è¯­æ°”åå¥½ å¹³å°åå¥½",
                top_k=10
            )
            
            preferences = {}
            for result in results:
                if result.memory and result.memory.metadata:
                    metadata = result.memory.metadata
                    ***REMOVED*** ä»metadataä¸­æå–åå¥½ä¿¡æ¯
                    for key in ["style_preference", "tone_preference", "platform_preference", "é£æ ¼åå¥½", "è¯­æ°”åå¥½", "å¹³å°åå¥½"]:
                        if key in metadata and metadata[key]:
                            pref_key = key.replace("_preference", "").replace("åå¥½", "")
                            if pref_key not in preferences:
                                preferences[pref_key] = metadata[key]
            
            return preferences
        except Exception as e:
            logger.debug(f"Failed to get auto-fill preferences: {e}")
            return {}
    
    def _get_auto_fill_shot_suggestions(self) -> List[str]:
        """
        ä»UniMemå†å²è®°å¿†ä¸­è·å–é•œå¤´ç´ æå»ºè®®
        
        Returns:
            é•œå¤´ç´ ææè¿°åˆ—è¡¨
        """
        if not self.unimem:
            return []
        
        try:
            ***REMOVED*** æ£€ç´¢å†å²è„šæœ¬ä¸­çš„é•œå¤´ä¿¡æ¯
            results = self.unimem.recall(
                query="é•œå¤´ç´ æ ç”»é¢æè¿° è§†é¢‘è„šæœ¬ é•œå¤´",
                top_k=5
            )
            
            shots = []
            for result in results:
                if result.memory:
                    ***REMOVED*** ä»metadataä¸­æå–é•œå¤´ä¿¡æ¯
                    metadata = result.memory.get("metadata", {}) if isinstance(result.memory, dict) else getattr(result.memory, "metadata", {})
                    if isinstance(metadata, dict):
                        script_segments = metadata.get("script_segments", [])
                        if script_segments:
                            for segment in script_segments[:3]:  ***REMOVED*** æ¯ä¸ªè„šæœ¬å–å‰3ä¸ªé•œå¤´
                                if isinstance(segment, dict):
                                    shot_desc = segment.get("ç”»é¢", segment.get("shot_description", ""))
                                    if shot_desc and shot_desc not in shots:
                                        shots.append(shot_desc)
            
            ***REMOVED*** å¦‚æœæ²¡æœ‰æ‰¾åˆ°ï¼Œè¿”å›é€šç”¨çš„é•œå¤´å»ºè®®
            if not shots:
                shots = [
                    "äº§å“æ•´ä½“å±•ç¤ºï¼ˆå…¨æ™¯ï¼Œå±•ç¤ºäº§å“å…¨è²Œï¼‰",
                    "äº§å“ç‰¹å†™ï¼ˆç»†èŠ‚å±•ç¤ºï¼Œçªå‡ºå–ç‚¹ï¼‰",
                    "ä½¿ç”¨åœºæ™¯å±•ç¤ºï¼ˆå®é™…åº”ç”¨åœºæ™¯ï¼‰",
                    "å¯¹æ¯”æ•ˆæœï¼ˆä½¿ç”¨å‰åå¯¹æ¯”ï¼‰",
                    "ç”¨æˆ·ååº”é•œå¤´ï¼ˆæƒŠå–œ/æ»¡æ„è¡¨æƒ…ï¼‰",
                    "äº§å“ç»†èŠ‚ç‰¹å†™ï¼ˆæè´¨/åšå·¥ç»†èŠ‚ï¼‰"
                ]
            
            return shots[:8]
        except Exception as e:
            logger.debug(f"Failed to get auto-fill shot suggestions: {e}")
            return []
    
    def parse_word_document(self, doc_path: str) -> Dict[str, Any]:
        """
        è§£æ Word æ–‡æ¡£ï¼Œæå–ç”¨æˆ·è®°å¿†ã€å•†å“ä¿¡æ¯ã€é•œå¤´ç´ æ
        
        ä½¿ç”¨ LLM API ä» Word æ–‡æ¡£ä¸­æå–çŸ­è§†é¢‘å‰§æœ¬ prompt ä¸­éœ€è¦çš„å‚æ•°ã€‚
        
        æ”¯æŒä¸‰ç§ç±»å‹çš„è®°å¿†æå–ï¼š
        1. å½“å‰ä»»åŠ¡è®°å¿†ï¼ˆéœ€æ±‚ï¼‰ï¼šå½“å‰è§†é¢‘åˆ›ä½œä»»åŠ¡çš„å…·ä½“éœ€æ±‚
        2. ä¿®æ”¹éœ€æ±‚è®°å¿†ï¼šå¯¹è¯ä¸­æå–çš„æ–°çš„ä¿®æ”¹éœ€æ±‚
        3. é€šç”¨è®°å¿†æ€»ç»“ï¼šè·¨ä»»åŠ¡çš„ç”¨æˆ·åå¥½å’Œé€šç”¨é£æ ¼åå¥½
        
        Args:
            doc_path: Word æ–‡æ¡£è·¯å¾„ï¼ˆ.docx æ ¼å¼ï¼‰
            
        Returns:
            è§£æç»“æœå­—å…¸ï¼ŒåŒ…å«ï¼š
            - task_memories: List[str] - å½“å‰ä»»åŠ¡è®°å¿†ï¼ˆéœ€æ±‚ï¼‰
            - modification_memories: List[str] - ä¿®æ”¹éœ€æ±‚è®°å¿†
            - general_memories: List[str] - è·¨ä»»åŠ¡é€šç”¨è®°å¿†æ€»ç»“
            - user_preferences: Dict[str, Any] - ç”¨æˆ·åå¥½ï¼ˆé£æ ¼ã€è¯­æ°”ã€å¹³å°åå¥½ç­‰ï¼‰
            - product_info: Dict[str, Any] - å•†å“ä¿¡æ¯
            - shot_materials: List[Dict[str, Any]] - é•œå¤´ç´ æä¿¡æ¯
            - video_type: str - è§†é¢‘ç±»å‹
            - platform: str - ç›®æ ‡å¹³å°
            - duration_seconds: int - ç›®æ ‡æ—¶é•¿ï¼ˆç§’ï¼‰
            
        Raises:
            AdapterNotAvailableError: å¦‚æœ python-docx ä¸å¯ç”¨
            AdapterError: å¦‚æœæ–‡æ¡£è·¯å¾„æ— æ•ˆæˆ–è§£æå¤±è´¥
        """
        if not DOCX_AVAILABLE:
            raise AdapterNotAvailableError(
                "python-docx library not available. Install with: pip install python-docx",
                adapter_name="VideoAdapter"
            )
        
        if not doc_path or not isinstance(doc_path, str):
            raise AdapterError("doc_path must be a non-empty string", adapter_name="VideoAdapter")
        
        doc_file = Path(doc_path)
        if not doc_file.exists():
            raise AdapterError(f"Document file not found: {doc_path}", adapter_name="VideoAdapter")
        
        if not doc_file.suffix.lower() == ".docx":
            raise AdapterError(
                f"Unsupported file format: {doc_file.suffix}. Only .docx files are supported",
                adapter_name="VideoAdapter"
            )
        
        try:
            doc = Document(doc_path)
            
            ***REMOVED*** è§£ææ–‡æ¡£å†…å®¹
            full_text = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    full_text.append(text)
            
            ***REMOVED*** å°†æ–‡æ¡£å†…å®¹åˆå¹¶ä¸ºæ–‡æœ¬
            document_text = "\n".join(full_text)
            
            ***REMOVED*** å¦‚æœæ–‡æ¡£ä¸ºç©ºï¼Œè¿”å›é»˜è®¤å€¼
            if not document_text.strip():
                logger.warning("Document is empty, returning default values")
                return {
                "task_memories": [],
                "modification_memories": [],
                "general_memories": [],
                "user_preferences": {},
                "product_info": {},
                "shot_materials": [],
                    "video_type": "ecommerce",
                    "platform": "douyin",
                    "duration_seconds": 60
            }
            
            ***REMOVED*** ä½¿ç”¨ LLM API æå–å‚æ•°
            try:
                prompt = f"""è¯·ä»ä»¥ä¸‹ Word æ–‡æ¡£å†…å®¹ä¸­æå–çŸ­è§†é¢‘å‰§æœ¬ç”Ÿæˆæ‰€éœ€çš„æ‰€æœ‰å‚æ•°ã€‚

æ–‡æ¡£å†…å®¹ï¼š
{document_text}

è¯·æå–ä»¥ä¸‹ä¿¡æ¯å¹¶ä»¥ JSON æ ¼å¼è¿”å›ï¼š

1. **è§†é¢‘åŸºæœ¬ä¿¡æ¯**ï¼ˆå¿…å¡«ï¼‰ï¼š
   - video_type: è§†é¢‘ç±»å‹ï¼Œä»ä»¥ä¸‹é€‰é¡¹é€‰æ‹©ï¼šecommerceï¼ˆç”µå•†æ¨å¹¿ï¼‰ã€ip_buildingï¼ˆä¸ªäººIPæ‰“é€ ï¼‰ã€knowledgeï¼ˆçŸ¥è¯†åˆ†äº«ï¼‰ã€vlogï¼ˆç”Ÿæ´»Vlogï¼‰ã€mediaï¼ˆè‡ªåª’ä½“å†…å®¹ï¼‰ã€‚å¦‚æœæœªæ‰¾åˆ°ï¼Œé»˜è®¤ä¸º "ecommerce"
   - platform: ç›®æ ‡å¹³å°ï¼Œä»ä»¥ä¸‹é€‰é¡¹é€‰æ‹©ï¼šdouyinï¼ˆæŠ–éŸ³ï¼‰ã€xiaohongshuï¼ˆå°çº¢ä¹¦ï¼‰ã€tiktokï¼ˆTikTokå›½é™…ï¼‰ã€youtubeï¼ˆYouTubeï¼‰ã€‚å¦‚æœæœªæ‰¾åˆ°ï¼Œé»˜è®¤ä¸º "douyin"
   - duration_seconds: ç›®æ ‡æ—¶é•¿ï¼ˆç§’ï¼‰ï¼Œæ•°å­—ã€‚å¦‚æœæœªæ‰¾åˆ°ï¼Œé»˜è®¤ä¸º 60

2. **å½“å‰ä»»åŠ¡éœ€æ±‚**ï¼ˆäºŒã€å½“å‰ä»»åŠ¡éœ€æ±‚ï¼‰ï¼š
   - task_memories: å­—ç¬¦ä¸²æ•°ç»„ï¼Œæ¯è¡Œä¸€æ¡éœ€æ±‚

3. **ä¿®æ”¹éœ€æ±‚**ï¼ˆä¸‰ã€ä¿®æ”¹éœ€æ±‚ï¼Œå¯é€‰ï¼‰ï¼š
   - modification_memories: å­—ç¬¦ä¸²æ•°ç»„ï¼Œæ¯è¡Œä¸€æ¡ä¿®æ”¹éœ€æ±‚

4. **é€šç”¨è®°å¿†æ€»ç»“**ï¼ˆå››ã€é€šç”¨è®°å¿†æ€»ç»“ï¼Œå¯é€‰ï¼‰ï¼š
   - general_memories: å­—ç¬¦ä¸²æ•°ç»„ï¼Œæ¯è¡Œä¸€æ¡é€šç”¨è®°å¿†

5. **ç”¨æˆ·åå¥½è®¾ç½®**ï¼ˆäº”ã€ç”¨æˆ·åå¥½è®¾ç½®ï¼Œå¯é€‰ï¼‰ï¼š
   - user_preferences: å¯¹è±¡ï¼Œé”®å€¼å¯¹æ ¼å¼ï¼ˆå¦‚ï¼šé£æ ¼åå¥½: çœŸè¯šè‡ªç„¶ï¼‰

6. **å•†å“ä¿¡æ¯**ï¼ˆå…­ã€å•†å“ä¿¡æ¯ï¼Œä»…ç”µå•†é¢˜æéœ€è¦ï¼Œå¯é€‰ï¼‰ï¼š
   - product_info: å¯¹è±¡ï¼Œé”®å€¼å¯¹æ ¼å¼ï¼ˆå¦‚ï¼šäº§å“åç§°: æ–°å“æ‰‹æœºï¼‰

7. **é•œå¤´ç´ æ**ï¼ˆä¸ƒã€é•œå¤´ç´ æï¼Œå¯é€‰ï¼‰ï¼š
   - shot_materials: å¯¹è±¡æ•°ç»„ï¼Œæ¯ä¸ªå¯¹è±¡åŒ…å« labelï¼ˆå¯é€‰ï¼‰å’Œ description

è¯·ä¸¥æ ¼æŒ‰ç…§ä»¥ä¸‹ JSON æ ¼å¼è¿”å›ï¼š
{{
    "video_type": "ecommerce",
    "platform": "douyin",
    "duration_seconds": 60,
    "task_memories": ["éœ€æ±‚1", "éœ€æ±‚2"],
    "modification_memories": ["ä¿®æ”¹éœ€æ±‚1"],
    "general_memories": ["é€šç”¨è®°å¿†1"],
    "user_preferences": {{
        "é£æ ¼åå¥½": "çœŸè¯šè‡ªç„¶",
        "è¯­æ°”åå¥½": "åƒæœ‹å‹åˆ†äº«"
    }},
    "product_info": {{
        "äº§å“åç§°": "æ–°å“æ‰‹æœº",
        "æ ¸å¿ƒå–ç‚¹": "æ€§ä»·æ¯”ã€æ‹ç…§"
    }},
    "shot_materials": [
        {{"label": "é•œå¤´1", "description": "äº§å“ç‰¹å†™ï¼Œå±•ç¤ºæ‰‹æœºå¤–è§‚"}},
        {{"description": "ä½¿ç”¨åœºæ™¯ï¼Œå¹´è½»äººæ‹ç…§"}}
    ]
}}

æ³¨æ„ï¼š
- å¦‚æœæŸä¸ªå­—æ®µåœ¨æ–‡æ¡£ä¸­ä¸å­˜åœ¨ï¼Œè¯·ä½¿ç”¨ç©ºæ•°ç»„ [] æˆ–ç©ºå¯¹è±¡ {{}} æˆ–é»˜è®¤å€¼
- è¿‡æ»¤æ‰æ¨¡æ¿æç¤ºå†…å®¹ï¼ˆå¦‚ "[è¯·å¡«å†™]"ã€"[ä¾‹å¦‚]"ã€"ä¾‹å¦‚ï¼š"ã€"ã€ç³»ç»Ÿå»ºè®®ã€‘"ã€"ã€æ‚¨ä¹Ÿå¯ä»¥å¡«å†™æˆ–ä¿®æ”¹ã€‘"ã€"âœ“" ç­‰æ ‡è®°ï¼‰
- é”®å€¼å¯¹æ ¼å¼æ”¯æŒå¤šç§æ ¼å¼ï¼š
  * ä¸­æ–‡å†’å·ï¼ˆï¼šï¼‰å’Œè‹±æ–‡å†’å·ï¼ˆ:ï¼‰
  * ç­‰å·ï¼ˆ=ï¼‰åˆ†éš”
  * ç©ºæ ¼åˆ†éš”çš„é”®å€¼å¯¹
- é•œå¤´ç´ ææ”¯æŒå¤šç§æ ¼å¼ï¼š
  * "é•œå¤´1: æè¿°"
  * "é•œå¤´1ï¼šæè¿°"
  * "é•œå¤´1=æè¿°"
  * "é•œå¤´1 æè¿°"
  * ç›´æ¥æè¿°ï¼ˆæ²¡æœ‰ç¼–å·ï¼‰
- ä»»åŠ¡è®°å¿†å’Œé€šç”¨è®°å¿†æ”¯æŒï¼š
  * æ¯è¡Œä¸€æ¡
  * åˆ—è¡¨æ ¼å¼ï¼ˆç”¨ - æˆ– * å¼€å¤´ï¼‰
  * ç¼–å·æ ¼å¼ï¼ˆ1. 2. 3.ï¼‰
- å¯¹äºæ ¼å¼ä¸è§„èŒƒçš„å¡«å†™ï¼Œå°½é‡æ™ºèƒ½æå–æœ‰ç”¨ä¿¡æ¯
- ç¡®ä¿æ‰€æœ‰å­—æ®µéƒ½å­˜åœ¨ï¼Œå³ä½¿å€¼ä¸ºç©º"""
                
                messages = [
                    {
                        "role": "system",
                        "content": """ä½ æ˜¯ä¸€ä¸ªä¸“ä¸šçš„æ–‡æ¡£è§£æåŠ©æ‰‹ï¼Œæ“…é•¿ä» Word æ–‡æ¡£ä¸­æå–ç»“æ„åŒ–ä¿¡æ¯ã€‚è¯·å§‹ç»ˆä»¥æœ‰æ•ˆçš„ JSON æ ¼å¼è¿”å›ç»“æœã€‚

é‡è¦èƒ½åŠ›ï¼š
1. å®¹é”™æ€§å¼ºï¼šå¯¹äºæ ¼å¼ä¸è§„èŒƒçš„å¡«å†™ï¼Œèƒ½å¤Ÿæ™ºèƒ½æå–æœ‰ç”¨ä¿¡æ¯
2. çµæ´»è§£æï¼šæ”¯æŒå¤šç§é”®å€¼å¯¹æ ¼å¼ï¼ˆå†’å·ã€ç­‰å·ã€ç©ºæ ¼ç­‰ï¼‰
3. è¿‡æ»¤å™ªéŸ³ï¼šè‡ªåŠ¨è¿‡æ»¤æ¨¡æ¿æç¤ºå†…å®¹ã€ç¤ºä¾‹å†…å®¹ã€ç³»ç»Ÿå»ºè®®æ ‡è®°ç­‰
4. æ™ºèƒ½æ¨æ–­ï¼šå¯¹äºä¸å®Œæ•´çš„ä¿¡æ¯ï¼Œèƒ½å¤Ÿæ ¹æ®ä¸Šä¸‹æ–‡åˆç†æ¨æ–­
5. ä¿è¯å®Œæ•´æ€§ï¼šç¡®ä¿è¿”å›çš„JSONåŒ…å«æ‰€æœ‰å¿…éœ€å­—æ®µ"""
                    },
                    {"role": "user", "content": prompt}
                ]
                
                _, response_text = ark_deepseek_v3_2(messages, max_new_tokens=2048)
                result = self._parse_json_response(response_text)
                
                if result:
                    ***REMOVED*** éªŒè¯å’Œè§„èŒƒåŒ–ç»“æœ
                    parsed_result = {
                        "task_memories": result.get("task_memories", []),
                        "modification_memories": result.get("modification_memories", []),
                        "general_memories": result.get("general_memories", []),
                        "user_preferences": result.get("user_preferences", {}),
                        "product_info": result.get("product_info", {}),
                        "shot_materials": result.get("shot_materials", []),
                        "video_type": result.get("video_type", "ecommerce"),
                        "platform": result.get("platform", "douyin"),
                        "duration_seconds": result.get("duration_seconds", 60)
                    }
                    
                    ***REMOVED*** éªŒè¯ video_type
                    valid_video_types = ["ecommerce", "ip_building", "knowledge", "vlog", "media"]
                    if parsed_result["video_type"] not in valid_video_types:
                        logger.warning(f"Invalid video_type '{parsed_result['video_type']}', using default 'ecommerce'")
                        parsed_result["video_type"] = "ecommerce"
                    
                    ***REMOVED*** éªŒè¯ platform
                    valid_platforms = ["douyin", "xiaohongshu", "tiktok", "youtube"]
                    if parsed_result["platform"] not in valid_platforms:
                        logger.warning(f"Invalid platform '{parsed_result['platform']}', using default 'douyin'")
                        parsed_result["platform"] = "douyin"
                    
                    ***REMOVED*** éªŒè¯ duration_seconds
                    try:
                        parsed_result["duration_seconds"] = int(parsed_result["duration_seconds"])
                        if parsed_result["duration_seconds"] <= 0:
                            parsed_result["duration_seconds"] = 60
                    except (ValueError, TypeError):
                        logger.warning(f"Invalid duration_seconds '{parsed_result['duration_seconds']}', using default 60")
                        parsed_result["duration_seconds"] = 60
                    
                    ***REMOVED*** ç¡®ä¿åˆ—è¡¨ç±»å‹
                    for key in ["task_memories", "modification_memories", "general_memories", "shot_materials"]:
                        if not isinstance(parsed_result[key], list):
                            parsed_result[key] = []
                    
                    ***REMOVED*** ç¡®ä¿å­—å…¸ç±»å‹
                    for key in ["user_preferences", "product_info"]:
                        if not isinstance(parsed_result[key], dict):
                            parsed_result[key] = {}
                    
                    logger.info(f"Parsed Word document using LLM: {len(parsed_result['task_memories'])} task memories, "
                               f"{len(parsed_result['modification_memories'])} modification memories, "
                               f"{len(parsed_result['general_memories'])} general memories, "
                               f"{len(parsed_result['user_preferences'])} preference fields, "
                               f"{len(parsed_result['product_info'])} product fields, "
                               f"{len(parsed_result['shot_materials'])} shots, "
                               f"video_type={parsed_result['video_type']}, platform={parsed_result['platform']}, "
                               f"duration={parsed_result['duration_seconds']}s")
                    
                    return parsed_result
                else:
                    logger.warning("LLM failed to parse document, falling back to default values")
                    raise ValueError("LLM parsing failed")
                    
            except Exception as llm_error:
                logger.warning(f"LLM parsing failed: {llm_error}, using default values")
                ***REMOVED*** å¦‚æœ LLM è§£æå¤±è´¥ï¼Œè¿”å›é»˜è®¤å€¼
                return {
                    "task_memories": [],
                    "modification_memories": [],
                    "general_memories": [],
                    "user_preferences": {},
                    "product_info": {},
                    "shot_materials": [],
                    "video_type": "ecommerce",
                    "platform": "douyin",
                    "duration_seconds": 60
                }
                
        except Exception as e:
            raise AdapterError(
                f"Failed to parse Word document: {e}",
                adapter_name="VideoAdapter",
                cause=e
            ) from e
    
    def generate_video_script(
        self,
        task_memories: Optional[List[str]] = None,
        modification_memories: Optional[List[str]] = None,
        general_memories: Optional[List[str]] = None,
        user_preferences: Optional[Dict[str, Any]] = None,
        product_info: Optional[Dict[str, Any]] = None,
        shot_materials: Optional[List[Dict[str, Any]]] = None,
        video_type: str = "ecommerce",
        duration_seconds: int = 60,
        platform: str = "douyin",
        script_len: Optional[int] = None,
        use_unimem_retrieval: bool = True,
        store_to_unimem: bool = True
    ) -> Dict[str, Any]:
        """
        ç”ŸæˆçŸ­è§†é¢‘æ–‡æ¡ˆå’Œå‰ªè¾‘è„šæœ¬
        
        æ”¯æŒå¤šç§é¢˜æï¼šç”µå•†æ¨å¹¿ã€ä¸ªäººIPæ‰“é€ ã€çŸ¥è¯†åˆ†äº«ã€ç”Ÿæ´» Vlog ç­‰
        é¢˜æç‰¹å®šéœ€æ±‚ï¼ˆå¦‚ç”µå•†å–ç‚¹ã€IPäººè®¾ã€çŸ¥è¯†ç»“æ„ç­‰ï¼‰åº”æ”¾åœ¨ç”¨æˆ·è®°å¿†çš„ Word æ–‡æ¡£ä¸­
        
        é€šç”¨åŠŸèƒ½ï¼ˆä¿ç•™åœ¨ prompt ä¸­ï¼‰ï¼š
        1. æ–‡æ¡ˆæ’°å†™ï¼šç”Ÿæˆå¸å¼•äººçš„æ–‡æ¡ˆï¼ˆ5-20å­—/å¥ï¼Œå¼€åœºå¸å¼•ã€ä¸­é—´é€’è¿›ã€ç»“å°¾è½¬åŒ–ï¼‰
        2. é•œå¤´åŒ¹é…ï¼šä¸ºæ¯ä¸ªæ–‡æ¡ˆæ®µè½åŒ¹é…æœ€ä½³é•œå¤´ï¼ˆä¸é‡å¤ä½¿ç”¨ï¼‰
        3. èŠ‚å¥æ§åˆ¶ï¼šæ§åˆ¶æ—¶é•¿å’ŒèŠ‚å¥
        4. å‰ªè¾‘è„šæœ¬ï¼šç”Ÿæˆå®Œæ•´çš„å‰ªè¾‘æŒ‡å¯¼
        
        æ·±åº¦ç»“åˆç”¨æˆ·è®°å¿†å’Œåå¥½ï¼š
        - ä»»åŠ¡è®°å¿†ï¼šå½“å‰ä»»åŠ¡çš„å…·ä½“éœ€æ±‚ï¼ˆåŒ…å«é¢˜æç‰¹å®šè¦æ±‚ï¼‰
        - ä¿®æ”¹è®°å¿†ï¼šå¯¹è¯ä¸­æå–çš„ä¿®æ”¹éœ€æ±‚
        - é€šç”¨è®°å¿†ï¼šè·¨ä»»åŠ¡çš„ç”¨æˆ·é£æ ¼åå¥½
        
        Args:
            task_memories: å½“å‰ä»»åŠ¡è®°å¿†ï¼ˆéœ€æ±‚ï¼‰åˆ—è¡¨ï¼Œåº”åŒ…å«é¢˜æç‰¹å®šéœ€æ±‚ï¼ˆå¦‚ç”µå•†å–ç‚¹ã€IPäººè®¾ç­‰ï¼‰
            modification_memories: ä¿®æ”¹éœ€æ±‚è®°å¿†åˆ—è¡¨
            general_memories: è·¨ä»»åŠ¡é€šç”¨è®°å¿†æ€»ç»“åˆ—è¡¨
            user_preferences: ç”¨æˆ·åå¥½å­—å…¸ï¼ˆé£æ ¼ã€è¯­æ°”ã€å¹³å°åå¥½ç­‰ï¼‰
            product_info: å•†å“ä¿¡æ¯å­—å…¸ï¼ˆä»…ç”µå•†é¢˜æéœ€è¦ï¼Œå…¶ä»–é¢˜æå¯ä¼ ç©ºï¼‰
            shot_materials: é•œå¤´ç´ æä¿¡æ¯åˆ—è¡¨ï¼ˆå€™é€‰é•œå¤´ï¼‰
            video_type: è§†é¢‘ç±»å‹ï¼Œ"ecommerce"ï¼ˆç”µå•†ï¼‰ã€"ip_building"ï¼ˆä¸ªäººIPï¼‰ã€"knowledge"ï¼ˆçŸ¥è¯†åˆ†äº«ï¼‰ã€"vlog"ï¼ˆç”Ÿæ´»Vlogï¼‰ç­‰
            duration_seconds: ç›®æ ‡è§†é¢‘æ—¶é•¿ï¼ˆç§’ï¼‰ï¼Œé»˜è®¤ 60 ç§’
            platform: å¹³å°ç±»å‹ï¼Œ"douyin"ï¼ˆæŠ–éŸ³ï¼‰ã€"xiaohongshu"ï¼ˆå°çº¢ä¹¦ï¼‰ã€"tiktok"ï¼ˆTikTokå›½é™…ï¼‰ã€"youtube"ï¼ˆYouTubeï¼‰
            script_len: è„šæœ¬æ®µæ•°ï¼ˆåºå·ä¸Šé™ï¼‰ï¼Œå¦‚æœä¸æä¾›åˆ™æ ¹æ®æ—¶é•¿è‡ªåŠ¨è®¡ç®—
            
        Returns:
            ç”Ÿæˆçš„è„šæœ¬å­—å…¸ï¼ˆJSON æ ¼å¼ï¼‰ï¼ŒåŒ…å«ï¼š
            - script: str - å®Œæ•´æ–‡æ¡ˆ
            - segments: List[Dict[str, Any]] - åˆ†æ®µè¯¦æƒ…ï¼ˆåŒ…å«åºå·ã€åˆ†ç±»ã€æ™¯åˆ«ã€æ–‡æ¡ˆã€ç”»é¢ï¼‰
            - editing_script: Dict[str, Any] - å‰ªè¾‘è„šæœ¬ï¼ˆåŒ…å«æ—¶é—´è½´ã€è½¬åœºç­‰ï¼‰
            - total_duration: float - é¢„è®¡æ€»æ—¶é•¿ï¼ˆç§’ï¼‰
            
        Raises:
            AdapterError: å¦‚æœå‚æ•°æ— æ•ˆ
        """
        if not self.is_available():
            logger.warning("VideoAdapter not available, cannot perform generate_video_script")
            return {}
        
        ***REMOVED*** æ”¯æŒå¤šç§è§†é¢‘ç±»å‹
        valid_video_types = ["ecommerce", "ip_building", "knowledge", "vlog", "media"]
        if video_type not in valid_video_types:
            raise AdapterError(
                f"Invalid video_type '{video_type}', must be one of: {', '.join(valid_video_types)}",
                adapter_name="VideoAdapter"
            )
        
        if duration_seconds <= 0:
            raise AdapterError(
                f"duration_seconds must be positive, got {duration_seconds}",
                adapter_name="VideoAdapter"
            )
        
        if platform not in ["douyin", "xiaohongshu", "tiktok", "youtube"]:
            raise AdapterError(
                f"Invalid platform '{platform}', must be one of: douyin, xiaohongshu, tiktok, youtube",
                adapter_name="VideoAdapter"
            )
        
        ***REMOVED*** æ ¹æ®æ—¶é•¿è‡ªåŠ¨è®¡ç®— script_lenï¼ˆå¦‚æœæœªæä¾›ï¼‰
        if script_len is None:
            ***REMOVED*** å¹³å‡æ¯æ®µ 8-12 ç§’
            script_len = max(5, min(30, duration_seconds // 8))
        
        if script_len <= 0:
            raise AdapterError(
                f"script_len must be positive, got {script_len}",
                adapter_name="VideoAdapter"
            )
        
        try:
            ***REMOVED*** ========== UniMem ä¼˜åŠ¿åˆ©ç”¨ ==========
            ***REMOVED*** 1. ä» UniMem ä¸­æ£€ç´¢ç›¸å…³å†å²è®°å¿†ï¼ˆå¦‚æœå¯ç”¨ï¼‰
            enriched_memories = {}
            if use_unimem_retrieval and self.unimem and task_memories:
                enriched_memories = self.enrich_memories_from_unimem(
                    task_memories=task_memories,
                    video_type=video_type,
                    top_k=10
                )
            
            ***REMOVED*** æ„å»ºä¸Šä¸‹æ–‡ä¿¡æ¯ï¼ˆæ·±åº¦èåˆç”¨æˆ·è®°å¿†å’Œ UniMem æ£€ç´¢ç»“æœï¼‰
            context_parts = []
            
            ***REMOVED*** 1. å½“å‰ä»»åŠ¡è®°å¿†ï¼ˆæœ€é«˜ä¼˜å…ˆçº§ï¼‰
            if task_memories:
                context_parts.append("ã€å½“å‰ä»»åŠ¡éœ€æ±‚ã€‘")
                for mem in task_memories[:15]:
                    context_parts.append(f"- {mem}")
            
            ***REMOVED*** 2. ä¿®æ”¹éœ€æ±‚è®°å¿†ï¼ˆé‡è¦ï¼Œéœ€è¦ä¼˜å…ˆåº”ç”¨ï¼‰
            if modification_memories:
                context_parts.append("\nã€æœ€æ–°ä¿®æ”¹éœ€æ±‚ã€‘ï¼ˆå¿…é¡»ä¸¥æ ¼æ‰§è¡Œï¼‰")
                for mem in modification_memories[:10]:
                    context_parts.append(f"- {mem}")
            
            ***REMOVED*** 3. é€šç”¨è®°å¿†æ€»ç»“ï¼ˆé£æ ¼å’Œåå¥½åŸºç¡€ï¼‰
            if general_memories:
                context_parts.append("\nã€ç”¨æˆ·é€šç”¨åå¥½å’Œé£æ ¼ã€‘")
                for mem in general_memories[:15]:
                    context_parts.append(f"- {mem}")
            
            ***REMOVED*** 4. UniMem æ£€ç´¢çš„å†å²ç›¸å…³åˆ›ä½œï¼ˆå¦‚æœæœ‰ï¼‰
            if enriched_memories.get("historical_scripts"):
                context_parts.append("\nã€å†å²ç›¸å…³åˆ›ä½œã€‘ï¼ˆæ¥è‡ª UniMem æ£€ç´¢ï¼‰")
                for script in enriched_memories["historical_scripts"][:5]:
                    context_parts.append(f"- [ç›¸ä¼¼åº¦: {script['score']:.2f}] {script['content']}")
            
            ***REMOVED*** 5. UniMem æ£€ç´¢çš„æˆåŠŸæ¨¡å¼ï¼ˆå¦‚æœæœ‰ï¼‰
            if enriched_memories.get("successful_patterns"):
                context_parts.append("\nã€æˆåŠŸåˆ›ä½œæ¨¡å¼ã€‘ï¼ˆæ¥è‡ª UniMem ç»éªŒï¼‰")
                for pattern in enriched_memories["successful_patterns"][:3]:
                    context_parts.append(f"- {pattern['pattern']}")
            
            ***REMOVED*** 6. UniMem æ£€ç´¢çš„ç”¨æˆ·é£æ ¼åå¥½ï¼ˆå¦‚æœæœ‰ï¼‰
            if enriched_memories.get("user_style_preferences"):
                context_parts.append("\nã€ç”¨æˆ·é£æ ¼åå¥½ã€‘ï¼ˆæ¥è‡ª UniMem å†å²è®°å¿†ï¼‰")
                for pref in enriched_memories["user_style_preferences"][:3]:
                    context_parts.append(f"- {pref['preference']}")
            
            ***REMOVED*** 7. ç”¨æˆ·åå¥½ï¼ˆç»“æ„åŒ–åå¥½ä¿¡æ¯ï¼‰
            if user_preferences:
                context_parts.append("\nã€ç”¨æˆ·åå¥½è®¾ç½®ã€‘")
                for key, value in list(user_preferences.items())[:10]:
                    if isinstance(value, list):
                        context_parts.append(f"- {key}: {', '.join(str(v) for v in value[:5])}")
                    else:
                        context_parts.append(f"- {key}: {value}")
            
            ***REMOVED*** 8. å•†å“ä¿¡æ¯
            if product_info:
                context_parts.append("\nå•†å“ä¿¡æ¯ï¼š")
                for key, value in list(product_info.items())[:10]:  ***REMOVED*** é™åˆ¶æ•°é‡
                    if isinstance(value, list):
                        context_parts.append(f"- {key}: {', '.join(value[:3])}")
                    else:
                        context_parts.append(f"- {key}: {value}")
            
            if shot_materials:
                context_parts.append("\nå¯ç”¨é•œå¤´ç´ æï¼š")
                for i, shot in enumerate(shot_materials[:15], 1):  ***REMOVED*** é™åˆ¶æ•°é‡
                    shot_desc = shot.get("description", shot.get("label", f"é•œå¤´{i}"))
                    context_parts.append(f"- é•œå¤´{i}: {shot_desc}")
            
            context_text = "\n".join(context_parts)
            
            ***REMOVED*** æ ¹æ®å¹³å°é€‰æ‹©é£æ ¼è¦æ±‚
            platform_styles = {
                "douyin": "èŠ‚å¥å¿«ã€æ¢—å¤šã€äº’åŠ¨æ€§å¼º",
                "xiaohongshu": "ç²¾è‡´æ„Ÿå¼ºã€ç”Ÿæ´»åŒ–ã€åˆ†äº«æ„Ÿå¼º",
                "tiktok": "å›½é™…åŒ–è¡¨è¾¾ã€ç®€æ´æœ‰åŠ›",
                "youtube": "ä¸“ä¸šæ€§å¼ºã€è¯¦ç»†è®²è§£"
            }
            
            platform_visuals = {
                "douyin": "å¿«èŠ‚å¥ã€è½¬åœºæµç•…ã€ç‰¹æ•ˆä¸°å¯Œ",
                "xiaohongshu": "ç”»é¢ç²¾è‡´ã€è‰²è°ƒç»Ÿä¸€ã€æ„å›¾è®²ç©¶",
                "tiktok": "å›½é™…åŒ–å®¡ç¾ã€ç®€æ´å¤§æ–¹",
                "youtube": "ç”»é¢ç¨³å®šã€ä¸“ä¸šæ„Ÿå¼º"
            }
            
            platform_style = platform_styles.get(platform, platform_styles["douyin"])
            platform_visual = platform_visuals.get(platform, platform_visuals["douyin"])
            
            ***REMOVED*** é€šç”¨ prompt æ¨¡æ¿ï¼ˆé¢˜æç‰¹å®šéœ€æ±‚ä»ç”¨æˆ·è®°å¿†è·å–ï¼‰
            ***REMOVED*** æ ¹æ®è§†é¢‘ç±»å‹æ·»åŠ åŸºç¡€è¯´æ˜
            video_type_descriptions = {
                "ecommerce": "ç”µå•†æ¨å¹¿è§†é¢‘ï¼ˆå–ç‚¹ã€è½¬åŒ–ç­‰éœ€æ±‚è¯·åœ¨ä»»åŠ¡è®°å¿†ä¸­æä¾›ï¼‰",
                "ip_building": "ä¸ªäººIPæ‰“é€ è§†é¢‘ï¼ˆäººè®¾ã€é£æ ¼ç­‰éœ€æ±‚è¯·åœ¨ä»»åŠ¡è®°å¿†ä¸­æä¾›ï¼‰",
                "knowledge": "çŸ¥è¯†åˆ†äº«è§†é¢‘ï¼ˆçŸ¥è¯†ç»“æ„ã€è®²è§£æ–¹å¼ç­‰éœ€æ±‚è¯·åœ¨ä»»åŠ¡è®°å¿†ä¸­æä¾›ï¼‰",
                "vlog": "ç”Ÿæ´» Vlog è§†é¢‘ï¼ˆä¸»é¢˜ã€æƒ…æ„Ÿè¡¨è¾¾ç­‰éœ€æ±‚è¯·åœ¨ä»»åŠ¡è®°å¿†ä¸­æä¾›ï¼‰",
                "media": "è‡ªåª’ä½“å†…å®¹è§†é¢‘ï¼ˆå†…å®¹æ–¹å‘ç­‰éœ€æ±‚è¯·åœ¨ä»»åŠ¡è®°å¿†ä¸­æä¾›ï¼‰"
            }
            
            video_type_desc = video_type_descriptions.get(video_type, "çŸ­è§†é¢‘å†…å®¹")
            
            ***REMOVED*** ç»Ÿä¸€çš„ promptï¼ˆé€šç”¨éœ€æ±‚ï¼Œé¢˜æç‰¹å®šéœ€æ±‚ä»ç”¨æˆ·è®°å¿†è·å–ï¼‰
            prompt = f"""è¯·æ ¹æ®ä»¥ä¸‹ä¿¡æ¯ç”Ÿæˆä¸€ä¸ªä¼˜è´¨çš„çŸ­è§†é¢‘æ–‡æ¡ˆå’Œå‰ªè¾‘è„šæœ¬ã€‚

**é‡è¦è¯´æ˜**ï¼š
- é¢˜æç±»å‹ï¼š{video_type_desc}
- é¢˜æç‰¹å®šçš„éœ€æ±‚ï¼ˆå¦‚ç”µå•†å–ç‚¹ã€IPäººè®¾ã€çŸ¥è¯†ç»“æ„ç­‰ï¼‰åº”ä»ä»»åŠ¡è®°å¿†ä¸­è·å–
- ä¸¥æ ¼éµå¾ªé€šç”¨çŸ­è§†é¢‘åˆ›ä½œè¦æ±‚
- ä»¥ JSON æ ¼å¼è¿”å›ï¼ˆJSON æ ¼å¼è´¨é‡æ›´å¥½ï¼‰

ç›®æ ‡ï¼š
- è§†é¢‘ç±»å‹ï¼š{video_type_desc}
- ç›®æ ‡æ—¶é•¿ï¼š{duration_seconds} ç§’
- å¹³å°ï¼š{platform}ï¼ˆ{platform_style}ï¼‰
- è„šæœ¬æ®µæ•°ï¼šä¸è¶…è¿‡ {script_len} æ®µ
- æ ¸å¿ƒç›®æ ‡ï¼šæ ¹æ®ä»»åŠ¡è®°å¿†ä¸­çš„å…·ä½“è¦æ±‚ç¡®å®šï¼ˆå¦‚å¸å¼•ç”¨æˆ·ã€çªå‡ºå–ç‚¹ã€å»ºç«‹äººè®¾ã€åˆ†äº«çŸ¥è¯†ã€è®°å½•ç”Ÿæ´»ç­‰ï¼‰

{context_text}

**é€šç”¨è¦æ±‚**ï¼ˆé€‚ç”¨äºæ‰€æœ‰é¢˜æï¼‰ï¼š

1. **æ–‡æ¡ˆè¦æ±‚**ï¼ˆä¸¥æ ¼éµå¾ªç”¨æˆ·è®°å¿†å’Œåå¥½ï¼‰ï¼š
   - **é•¿åº¦**ï¼šæ¯å¥5-20å­—ï¼Œé‡ç‚¹çªå‡º
   - **å¼€åœºï¼ˆé»„é‡‘3ç§’ï¼‰**ï¼šå¿…é¡»å¸å¼•äººï¼Œç”¨æ‚¬å¿µæˆ–å…±é¸£ç‚¹æŠ“ä½è§‚ä¼—
     - ç»“åˆç”¨æˆ·é€šç”¨åå¥½ï¼Œä½¿ç”¨ç”¨æˆ·å–œæ¬¢çš„è¡¨è¾¾æ–¹å¼
     - æ ¹æ®ä»»åŠ¡è®°å¿†ä¸­çš„é¢˜æç‰¹å®šè¦æ±‚è°ƒæ•´
   - **ä¸­é—´**ï¼šå±‚å±‚é€’è¿›ï¼Œçªå‡ºæ ¸å¿ƒå†…å®¹
     - æ ¹æ®ä»»åŠ¡è®°å¿†ä¸­çš„å…·ä½“è¦æ±‚ï¼ˆå¦‚ç”µå•†å–ç‚¹ã€IPäººè®¾ã€çŸ¥è¯†è¦ç‚¹ã€Vlogä¸»é¢˜ç­‰ï¼‰å±•å¼€
     - ç»“åˆç”¨æˆ·åå¥½ï¼Œçªå‡ºç”¨æˆ·å…³å¿ƒçš„å†…å®¹
     - æ ¹æ®ä¿®æ”¹éœ€æ±‚è®°å¿†ï¼Œç¡®ä¿æ»¡è¶³æœ€æ–°è¦æ±‚
   - **ç»“å°¾**ï¼šæ ¹æ®é¢˜æç±»å‹ç¡®å®šï¼ˆå¦‚è½¬åŒ–ã€äº’åŠ¨ã€æ€»ç»“ç­‰ï¼Œå…·ä½“è§ä»»åŠ¡è®°å¿†ï¼‰
   - **è¯­æ°”å’Œé£æ ¼**ï¼ˆé‡ç‚¹ç»“åˆç”¨æˆ·è®°å¿†ï¼‰ï¼š
     - æ ¹æ®é€šç”¨è®°å¿†å’Œç”¨æˆ·åå¥½ï¼Œä½¿ç”¨ç”¨æˆ·å–œæ¬¢çš„è¯­æ°”
     - æ ¹æ®ä»»åŠ¡è®°å¿†ä¸­çš„é¢˜æè¦æ±‚è°ƒæ•´ï¼ˆå¦‚ç”µå•†çš„ç”Ÿæ´»åŒ–åˆ†äº«ã€IPçš„ä¸“ä¸šæ„Ÿã€çŸ¥è¯†çš„æ˜“æ‡‚æ€§ã€Vlogçš„çœŸå®æ„Ÿç­‰ï¼‰
     - é¿å…ç”Ÿç¡¬çš„è¡¨è¾¾ï¼Œç”¨ç”Ÿæ´»åŒ–ã€å£è¯­åŒ–çš„è¡¨è¾¾
     - é€‚å½“ä½¿ç”¨æµè¡Œè¯­å’Œç½‘ç»œçƒ­è¯ï¼Œä½†ä¸è¦è¿‡åº¦
     - å¯ä»¥åŠ å…¥ä¸ªäººæ„Ÿå—å’ŒçœŸå®ä½“éªŒï¼ˆå‚è€ƒé€šç”¨è®°å¿†ä¸­çš„ä½“éªŒï¼‰
   - **ç¦æ­¢å†…å®¹**ï¼š
     - é¿å…"å§å¦¹ä»¬"ã€"å°å§å§"ç­‰ç§°å‘¼ï¼ˆé™¤éç”¨æˆ·åå¥½ä¸­æœ‰æ˜ç¡®è¦æ±‚ï¼‰
     - é¿å…è´Ÿé¢å†…å®¹å’ŒåŠ¨ç‰©æ€§åˆ¶å“(è‚‰è›‹å¥¶ç­‰)å’Œäº”è¾›
   - **å¹³å°é£æ ¼**ï¼ˆ{platform}ï¼‰ï¼š
     - {platform_style}

2. **ç”»é¢è¦æ±‚**ï¼ˆé€šç”¨åŒ¹é…è§„åˆ™ï¼‰ï¼š
   - **ç”»é¢æè¿°**ï¼šä¸è¦é‡å¤ä½¿ç”¨ï¼Œè¦å’Œå€™é€‰é•œå¤´ç”»é¢æè¿°ä¸€è‡´
   - **ç²¾å‡†åŒ¹é…**ï¼šç”»é¢å’Œæ–‡æ¡ˆè¦ç²¾å‡†åŒ¹é…
   - **æ™¯åˆ«é€‰é¡¹**ï¼šè¿œæ™¯ã€å…¨æ™¯ã€ä¸­æ™¯ã€è¿‘æ™¯ã€ç‰¹å†™ã€ä¿¯æ‹ã€ä»°æ‹ã€è·Ÿè¸ªé•œå¤´
   - **æ™¯åˆ«åŒ¹é…**ï¼šæ™¯åˆ«è¦å’Œç”»é¢å†…å®¹ç›¸ç¬¦
   - **é•œå¤´åˆ‡æ¢**ï¼šè¦æµç•…è‡ªç„¶
   - **é‡ç‚¹ç”»é¢**ï¼šé‡ç‚¹ç”»é¢è¦ç»™ç‰¹å†™
   - **å¹³å°é£æ ¼**ï¼ˆ{platform}ï¼‰ï¼š
     - {platform_visual}

3. **é•œå¤´ä½¿ç”¨**ï¼š
   - å€™é€‰é•œå¤´ä¸è¦é‡å¤ä½¿ç”¨
   - åºå·ä¸è¶…è¿‡ {script_len}

è¯·ä»¥ JSON æ ¼å¼è¿”å›ç»“æœï¼š
{{
    "script": "å®Œæ•´æ–‡æ¡ˆï¼ˆåŒ…å«åˆ†æ®µï¼‰",
    "segments": [
        {{
            "åºå·": 1,
            "åˆ†ç±»": "åˆ†ç±»é€‰é¡¹ï¼ˆæ ¹æ®å®é™…æƒ…å†µå’Œé¢˜æç±»å‹ï¼‰",
            "æ™¯åˆ«": "æ™¯åˆ«é€‰é¡¹ï¼ˆè¿œæ™¯/å…¨æ™¯/ä¸­æ™¯/è¿‘æ™¯/ç‰¹å†™/ä¿¯æ‹/ä»°æ‹/è·Ÿè¸ªé•œå¤´ï¼‰",
            "æ–‡æ¡ˆ": "æ–‡æ¡ˆå†…å®¹ï¼ˆ5-20å­—/å¥ï¼Œæ ¹æ®ä»»åŠ¡è®°å¿†ä¸­çš„é¢˜æè¦æ±‚ï¼‰",
            "ç”»é¢": "ç”»é¢æè¿°ï¼ˆä¸å€™é€‰é•œå¤´ä¸€è‡´ï¼Œä¸é‡å¤ï¼‰",
            "duration_seconds": 8,
            "matched_shots": ["é•œå¤´1", "é•œå¤´2"],
            "shot_reason": "ä¸ºä»€ä¹ˆé€‰æ‹©è¿™äº›é•œå¤´",
            "visual_elements": ["ç‰¹æ•ˆå»ºè®®", "å­—å¹•æ—¶æœº"],
            "music_cue": "éŸ³ä¹èŠ‚å¥å»ºè®®"
        }},
        ...
    ],
    "editing_script": {{
        "total_duration": {duration_seconds},
        "transitions": [
            {{"from_segment": 1, "to_segment": 2, "type": "è½¬åœºæ–¹å¼"}},
            ...
        ],
        "music_suggestions": {{
            "style": "éŸ³ä¹é£æ ¼å»ºè®®",
            "key_moments": ["å…³é”®æ—¶åˆ»1", "å…³é”®æ—¶åˆ»2"]
        }},
        "special_effects": [
            {{"segment": 1, "effect": "ç‰¹æ•ˆå»ºè®®", "timing": "æ—¶æœº"}},
            ...
        ]
    }},
    "summary": {{
        "hook": "å¼€å¤´äº®ç‚¹æ€»ç»“",
        "core_content": "æ ¸å¿ƒå†…å®¹æ€»ç»“ï¼ˆæ ¹æ®é¢˜æç±»å‹ï¼šå–ç‚¹/IPäººè®¾/çŸ¥è¯†ç‚¹/ç”Ÿæ´»ä¸»é¢˜ç­‰ï¼‰",
        "ending_type": "ç»“å°¾ç±»å‹ï¼ˆè½¬åŒ–/äº’åŠ¨/æ€»ç»“ç­‰ï¼Œæ ¹æ®ä»»åŠ¡è®°å¿†ï¼‰",
        "target_audience": "ç›®æ ‡å—ä¼—",
        "style_applied": "åº”ç”¨çš„é£æ ¼ç‰¹ç‚¹ï¼ˆåŸºäºç”¨æˆ·è®°å¿†ï¼‰"
    }}
}}"""
            
            messages = [
                {
                    "role": "system",
                    "content": "ä½ æ˜¯ä¸€ä¸ªä¸“ä¸šçš„çŸ­è§†é¢‘æ–‡æ¡ˆå’Œå‰ªè¾‘æŒ‡å¯¼ä¸“å®¶ï¼Œæ“…é•¿ç”Ÿæˆå¸å¼•äººçš„æ–‡æ¡ˆå’Œç²¾å‡†çš„å‰ªè¾‘è„šæœ¬ã€‚è¯·å§‹ç»ˆä»¥æœ‰æ•ˆçš„ JSON æ ¼å¼è¿”å›ç»“æœã€‚"
                },
                {"role": "user", "content": prompt}
            ]
            
            _, response_text = ark_deepseek_v3_2(messages, max_new_tokens=4096)
            result = self._parse_json_response(response_text)
            
            if result:
                logger.info(f"Generated video script: {len(result.get('segments', []))} segments, "
                           f"target duration: {duration_seconds}s")
                
                ***REMOVED*** ========== UniMem ä¼˜åŠ¿åˆ©ç”¨ï¼šå­˜å‚¨ç”Ÿæˆçš„è„šæœ¬ ==========
                if store_to_unimem and self.unimem:
                    memory_id = self.store_script_to_unimem(
                        script_data=result,
                        task_memories=task_memories or [],
                        video_type=video_type,
                        platform=platform
                    )
                    if memory_id:
                        result["unimem_memory_id"] = memory_id
                        logger.info(f"Script stored to UniMem with memory_id: {memory_id}")
                
                return result
            else:
                logger.warning(f"Failed to parse video script response: {response_text[:200]}")
                return {}
        except Exception as e:
            logger.error(f"Error generating video script: {e}", exc_info=True)
            return {}
    
    def match_shots_to_script(
        self,
        script_segments: List[Dict[str, Any]],
        available_shots: List[Dict[str, Any]]
    ) -> List[Dict[str, Any]]:
        """
        ä¸ºæ–‡æ¡ˆæ®µè½åŒ¹é…æœ€ä½³é•œå¤´ç´ æ
        
        ä½¿ç”¨è¯­ä¹‰ç›¸ä¼¼åº¦åŒ¹é…æ–‡æ¡ˆå’Œé•œå¤´ï¼Œä¸ºæ¯ä¸ªæ®µè½æ¨èæœ€åˆé€‚çš„é•œå¤´ã€‚
        
        Args:
            script_segments: æ–‡æ¡ˆæ®µè½åˆ—è¡¨ï¼Œæ¯ä¸ªæ®µè½åŒ…å« script_text
            available_shots: å¯ç”¨é•œå¤´åˆ—è¡¨ï¼Œæ¯ä¸ªé•œå¤´åŒ…å« description æˆ– label
            
        Returns:
            åŒ¹é…ç»“æœåˆ—è¡¨ï¼Œæ¯ä¸ªç»“æœåŒ…å«ï¼š
            - segment: åŸæ®µè½ä¿¡æ¯
            - recommended_shots: æ¨èçš„é•œå¤´åˆ—è¡¨ï¼ˆæŒ‰åŒ¹é…åº¦æ’åºï¼‰
            - match_scores: åŒ¹é…åˆ†æ•°åˆ—è¡¨
            
        Raises:
            AdapterError: å¦‚æœå‚æ•°æ— æ•ˆ
        """
        if not self.is_available():
            logger.warning("VideoAdapter not available, cannot perform match_shots_to_script")
            return []
        
        if not script_segments or not isinstance(script_segments, list):
            raise AdapterError("script_segments must be a non-empty list", adapter_name="VideoAdapter")
        
        if not available_shots or not isinstance(available_shots, list):
            raise AdapterError("available_shots must be a non-empty list", adapter_name="VideoAdapter")
        
        try:
            results = []
            
            for segment in script_segments:
                script_text = segment.get("script_text", "")
                if not script_text:
                    continue
                
                ***REMOVED*** ä½¿ç”¨è¯­ä¹‰æ£€ç´¢åŒ¹é…é•œå¤´
                ***REMOVED*** å°†é•œå¤´æè¿°è½¬æ¢ä¸ºæŸ¥è¯¢åˆ—è¡¨
                shot_descriptions = []
                for shot in available_shots:
                    desc = shot.get("description", shot.get("label", ""))
                    if desc:
                        shot_descriptions.append(desc)
                
                if not shot_descriptions:
                    results.append({
                        "segment": segment,
                        "recommended_shots": [],
                        "match_scores": []
                    })
                    continue
                
                ***REMOVED*** ä½¿ç”¨è¯­ä¹‰æ£€ç´¢æ‰¾åˆ°æœ€ç›¸å…³çš„é•œå¤´
                ***REMOVED*** ç®€åŒ–å®ç°ï¼šä½¿ç”¨å…³é”®è¯åŒ¹é…å’Œç›¸ä¼¼åº¦è®¡ç®—
                recommended_shots = []
                match_scores = []
                
                for shot in available_shots:
                    shot_desc = shot.get("description", shot.get("label", ""))
                    if not shot_desc:
                        continue
                    
                    ***REMOVED*** ç®€å•çš„ç›¸ä¼¼åº¦è®¡ç®—ï¼ˆå¯ä»¥åç»­ä¼˜åŒ–ä¸ºå‘é‡ç›¸ä¼¼åº¦ï¼‰
                    score = self._calculate_text_similarity(script_text, shot_desc)
                    recommended_shots.append(shot)
                    match_scores.append(score)
                
                ***REMOVED*** æŒ‰åˆ†æ•°æ’åº
                sorted_pairs = sorted(
                    zip(recommended_shots, match_scores),
                    key=lambda x: x[1],
                    reverse=True
                )
                
                recommended_shots = [shot for shot, _ in sorted_pairs[:5]]  ***REMOVED*** å–å‰5ä¸ª
                match_scores = [score for _, score in sorted_pairs[:5]]
                
                results.append({
                    "segment": segment,
                    "recommended_shots": recommended_shots,
                    "match_scores": match_scores
                })
            
            logger.debug(f"Matched shots for {len(results)} script segments")
            return results
        except Exception as e:
            logger.error(f"Error matching shots to script: {e}", exc_info=True)
            return []
    
    def extract_modification_memories_from_conversation(
        self,
        conversation_text: str,
        existing_modifications: Optional[List[str]] = None
    ) -> List[str]:
        """
        ä»å¯¹è¯ä¸­æå–æ–°çš„ä¿®æ”¹éœ€æ±‚è®°å¿†
        
        åˆ†æå¯¹è¯å†…å®¹ï¼Œè¯†åˆ«ç”¨æˆ·æå‡ºçš„ä¿®æ”¹è¦æ±‚ã€è°ƒæ•´æ„è§ç­‰ï¼Œ
        å¹¶æå–ä¸ºç»“æ„åŒ–çš„ä¿®æ”¹éœ€æ±‚è®°å¿†ã€‚
        
        æ”¯æŒå¤šæ¬¡å¯¹è¯ç´¯ç§¯ä¿®æ”¹éœ€æ±‚ï¼šå¦‚æœæä¾›äº†å·²æœ‰çš„ä¿®æ”¹éœ€æ±‚ï¼Œä¼šé¿å…é‡å¤æå–ã€‚
        
        Args:
            conversation_text: å¯¹è¯æ–‡æœ¬
            existing_modifications: å·²æœ‰çš„ä¿®æ”¹éœ€æ±‚åˆ—è¡¨ï¼ˆå¯é€‰ï¼‰ï¼Œç”¨äºé¿å…é‡å¤æå–
            
        Returns:
            ä¿®æ”¹éœ€æ±‚è®°å¿†åˆ—è¡¨ï¼Œæ¯ä¸ªå…ƒç´ æ˜¯ä¸€ä¸ªå…·ä½“çš„ä¿®æ”¹éœ€æ±‚ï¼ˆåªåŒ…å«æ–°å¢çš„ï¼Œä¸åŒ…å«å·²æœ‰çš„ï¼‰
            
        Raises:
            AdapterError: å¦‚æœå‚æ•°æ— æ•ˆ
        """
        if not conversation_text or not isinstance(conversation_text, str) or not conversation_text.strip():
            raise AdapterError("conversation_text cannot be empty", adapter_name="VideoAdapter")
        
        if not self.is_available():
            logger.warning("VideoAdapter not available, cannot extract modification memories")
            return []
        
        try:
            ***REMOVED*** æ„å»ºå·²æœ‰ä¿®æ”¹éœ€æ±‚çš„ä¸Šä¸‹æ–‡ï¼ˆå¦‚æœæœ‰ï¼‰
            existing_context = ""
            if existing_modifications:
                existing_context = f"""

å·²æœ‰ä¿®æ”¹éœ€æ±‚ï¼ˆè¯·é¿å…é‡å¤æå–ï¼‰ï¼š
{chr(10).join(f"- {mod}" for mod in existing_modifications[:10])}
{f"... è¿˜æœ‰ {len(existing_modifications) - 10} æ¡ä¿®æ”¹éœ€æ±‚" if len(existing_modifications) > 10 else ""}

è¯·åªæå–æœ¬æ¬¡å¯¹è¯ä¸­æ–°å¢çš„ã€ä¸å·²æœ‰ä¿®æ”¹éœ€æ±‚ä¸åŒçš„ä¿®æ”¹éœ€æ±‚ã€‚å¦‚æœç”¨æˆ·åªæ˜¯é‡å¤æˆ–å¼ºè°ƒå·²æœ‰éœ€æ±‚ï¼Œè¯·ä¸è¦é‡å¤æå–ã€‚"""
            
            prompt = f"""è¯·ä»ä»¥ä¸‹å¯¹è¯ä¸­æå–ç”¨æˆ·æå‡ºçš„**æ–°çš„**ä¿®æ”¹éœ€æ±‚ã€è°ƒæ•´æ„è§å’Œè¦æ±‚ã€‚

å¯¹è¯å†…å®¹ï¼š
{conversation_text[:2000]}{existing_context}

è¯·è¯†åˆ«å¹¶æå–ï¼š
1. æ–‡æ¡ˆç›¸å…³çš„ä¿®æ”¹è¦æ±‚ï¼ˆè¯­æ°”ã€é£æ ¼ã€é•¿åº¦ã€å†…å®¹ç­‰ï¼‰
2. ç”»é¢ç›¸å…³çš„ä¿®æ”¹è¦æ±‚ï¼ˆæ™¯åˆ«ã€é•œå¤´ã€è§†è§‰æ•ˆæœç­‰ï¼‰
3. èŠ‚å¥å’Œæ—¶é•¿ç›¸å…³çš„è°ƒæ•´
4. å•†å“ä¿¡æ¯ç›¸å…³çš„ä¿®æ”¹
5. å…¶ä»–æ˜ç¡®çš„ä¿®æ”¹æŒ‡ä»¤

é‡è¦æç¤ºï¼š
- åªæå–æœ¬æ¬¡å¯¹è¯ä¸­**æ–°å¢çš„**ä¿®æ”¹éœ€æ±‚
- å¦‚æœç”¨æˆ·åªæ˜¯é‡å¤æˆ–å¼ºè°ƒå·²æœ‰éœ€æ±‚ï¼Œè¯·ä¸è¦é‡å¤æå–
- å¦‚æœç”¨æˆ·ä¿®æ”¹æˆ–æ›´æ–°äº†å·²æœ‰éœ€æ±‚ï¼Œè¯·æå–æ›´æ–°åçš„éœ€æ±‚ï¼ˆå¯ä»¥è¦†ç›–åŸæœ‰éœ€æ±‚ï¼‰

è¯·ä»¥ JSON æ ¼å¼è¿”å›ç»“æœï¼š
{{
    "modification_memories": [
        "ä¿®æ”¹éœ€æ±‚1ï¼ˆå…·ä½“ã€æ˜ç¡®ï¼‰",
        "ä¿®æ”¹éœ€æ±‚2ï¼ˆå…·ä½“ã€æ˜ç¡®ï¼‰",
        ...
    ],
    "priority": "high/medium/low",  // ä¿®æ”¹çš„ç´§æ€¥ç¨‹åº¦
    "is_update": false  // æ˜¯å¦æ˜¯æ›´æ–°å·²æœ‰éœ€æ±‚
}}"""
            
            messages = [
                {
                    "role": "system",
                    "content": "ä½ æ˜¯ä¸€ä¸ªä¸“ä¸šçš„éœ€æ±‚åˆ†æåŠ©æ‰‹ï¼Œæ“…é•¿ä»å¯¹è¯ä¸­æå–æ˜ç¡®çš„ä¿®æ”¹éœ€æ±‚å’Œè°ƒæ•´æ„è§ã€‚è¯·å§‹ç»ˆä»¥æœ‰æ•ˆçš„ JSON æ ¼å¼è¿”å›ç»“æœï¼Œå¹¶é¿å…é‡å¤æå–å·²æœ‰éœ€æ±‚ã€‚"
                },
                {"role": "user", "content": prompt}
            ]
            
            _, response_text = ark_deepseek_v3_2(messages, max_new_tokens=1024)
            result = self._parse_json_response(response_text)
            
            if result and "modification_memories" in result:
                memories = result["modification_memories"]
                priority = result.get("priority", "medium")
                is_update = result.get("is_update", False)
                logger.info(f"Extracted {len(memories)} new modification memories "
                           f"(priority: {priority}, is_update: {is_update})")
                return memories
            else:
                logger.warning("Failed to extract modification memories from conversation")
                return []
        except Exception as e:
            logger.error(f"Error extracting modification memories: {e}", exc_info=True)
            return []
    
    def link_general_memories(
        self,
        task_memories: List[str],
        existing_general_memories: Optional[List[str]] = None
    ) -> List[str]:
        """
        è”ç³»è·¨ä»»åŠ¡çš„é€šç”¨è®°å¿†æ€»ç»“
        
        æ ¹æ®å½“å‰ä»»åŠ¡è®°å¿†ï¼Œä»å·²æœ‰çš„é€šç”¨è®°å¿†ä¸­ç­›é€‰å‡ºç›¸å…³çš„é€šç”¨åå¥½å’Œé£æ ¼æ€»ç»“ï¼Œ
        å»ºç«‹ä»»åŠ¡è®°å¿†ä¸é€šç”¨è®°å¿†ä¹‹é—´çš„è”ç³»ã€‚
        
        Args:
            task_memories: å½“å‰ä»»åŠ¡è®°å¿†åˆ—è¡¨
            existing_general_memories: å·²æœ‰çš„é€šç”¨è®°å¿†åˆ—è¡¨ï¼ˆå¯é€‰ï¼‰
            
        Returns:
            ä¸å½“å‰ä»»åŠ¡ç›¸å…³çš„é€šç”¨è®°å¿†åˆ—è¡¨
            
        Raises:
            AdapterError: å¦‚æœå‚æ•°æ— æ•ˆ
        """
        if not task_memories or not isinstance(task_memories, list):
            raise AdapterError("task_memories must be a non-empty list", adapter_name="VideoAdapter")
        
        if not existing_general_memories:
            return []
        
        if not self.is_available():
            logger.warning("VideoAdapter not available, cannot link general memories")
            return existing_general_memories[:10]  ***REMOVED*** è¿”å›å‰10ä¸ªä½œä¸ºé»˜è®¤
        
        try:
            ***REMOVED*** ä½¿ç”¨è¯­ä¹‰ç›¸ä¼¼åº¦ç­›é€‰ç›¸å…³çš„é€šç”¨è®°å¿†
            task_text = " ".join(task_memories[:5])  ***REMOVED*** ä½¿ç”¨å‰5ä¸ªä»»åŠ¡è®°å¿†ä½œä¸ºæŸ¥è¯¢
            
            relevant_memories = []
            for gen_mem in existing_general_memories:
                ***REMOVED*** ç®€å•çš„å…³é”®è¯åŒ¹é…ï¼ˆå¯ä»¥åç»­ä¼˜åŒ–ä¸ºå‘é‡ç›¸ä¼¼åº¦ï¼‰
                similarity = self._calculate_text_similarity(task_text, gen_mem)
                if similarity > 0.1:  ***REMOVED*** é˜ˆå€¼å¯ä»¥è°ƒæ•´
                    relevant_memories.append((gen_mem, similarity))
            
            ***REMOVED*** æŒ‰ç›¸ä¼¼åº¦æ’åºï¼Œå–å‰15ä¸ª
            relevant_memories.sort(key=lambda x: x[1], reverse=True)
            result = [mem for mem, _ in relevant_memories[:15]]
            
            logger.debug(f"Linked {len(result)} relevant general memories from {len(existing_general_memories)} total")
            return result
        except Exception as e:
            logger.error(f"Error linking general memories: {e}", exc_info=True)
            return existing_general_memories[:10]  ***REMOVED*** é™çº§è¿”å›
    
    def _calculate_text_similarity(self, text1: str, text2: str) -> float:
        """
        è®¡ç®—ä¸¤ä¸ªæ–‡æœ¬çš„ç›¸ä¼¼åº¦ï¼ˆç®€åŒ–å®ç°ï¼‰
        
        Args:
            text1: æ–‡æœ¬1
            text2: æ–‡æœ¬2
            
        Returns:
            ç›¸ä¼¼åº¦åˆ†æ•°ï¼ˆ0-1ï¼‰
        """
        if not text1 or not text2:
            return 0.0
        
        ***REMOVED*** ç®€å•çš„å…³é”®è¯åŒ¹é…ç›¸ä¼¼åº¦
        words1 = set(text1.lower().split())
        words2 = set(text2.lower().split())
        
        if not words1 or not words2:
            return 0.0
        
        intersection = words1.intersection(words2)
        union = words1.union(words2)
        
        if not union:
            return 0.0
        
        return len(intersection) / len(union)
    
    def optimize_script_for_editing(
        self,
        script_data: Dict[str, Any],
        feedback: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        ä¼˜åŒ–è„šæœ¬ä»¥æ›´å¥½åœ°é€‚é…å‰ªè¾‘æµç¨‹
        
        æ ¹æ®å‰ªè¾‘åé¦ˆæˆ–æœ€ä½³å®è·µï¼Œä¼˜åŒ–æ–‡æ¡ˆå’Œé•œå¤´åŒ¹é…ã€‚
        
        Args:
            script_data: åŸå§‹è„šæœ¬æ•°æ®
            feedback: å‰ªè¾‘åé¦ˆï¼ˆå¯é€‰ï¼‰
            
        Returns:
            ä¼˜åŒ–åçš„è„šæœ¬æ•°æ®ï¼Œç»“æ„ä¸åŸå§‹æ•°æ®ç›¸åŒ
        """
        if not self.is_available():
            logger.warning("VideoAdapter not available, cannot perform optimize_script_for_editing")
            return script_data
        
        if not script_data or not isinstance(script_data, dict):
            raise AdapterError("script_data must be a non-empty dict", adapter_name="VideoAdapter")
        
        try:
            ***REMOVED*** ä½¿ç”¨ LLM ä¼˜åŒ–è„šæœ¬
            prompt = f"""è¯·ä¼˜åŒ–ä»¥ä¸‹çŸ­è§†é¢‘è„šæœ¬ï¼Œä½¿å…¶æ›´ç¬¦åˆå‰ªè¾‘æµç¨‹å’Œæœ€ä½³å®è·µã€‚

åŸå§‹è„šæœ¬ï¼š
{json.dumps(script_data, ensure_ascii=False, indent=2)}

{f"å‰ªè¾‘åé¦ˆï¼š{feedback}" if feedback else ""}

ä¼˜åŒ–è¦æ±‚ï¼š
1. ç¡®ä¿æ¯æ®µæ–‡æ¡ˆæ—¶é•¿åˆç†ï¼ˆ5-15ç§’ï¼‰
2. ä¼˜åŒ–é•œå¤´åŒ¹é…ï¼Œç¡®ä¿è§†è§‰ä¸æ–‡æ¡ˆé«˜åº¦å¥‘åˆ
3. è°ƒæ•´èŠ‚å¥ï¼Œç¡®ä¿æ•´ä½“æµç•…
4. ä¼˜åŒ–è½¬åœºå’Œç‰¹æ•ˆå»ºè®®
5. ç¡®ä¿é»„é‡‘3ç§’å¼€å¤´è¶³å¤Ÿå¸å¼•äºº

è¯·è¿”å›ä¼˜åŒ–åçš„å®Œæ•´è„šæœ¬ï¼ˆJSON æ ¼å¼ï¼‰ï¼Œä¿æŒåŸæœ‰ç»“æ„ï¼š"""
            
            messages = [
                {
                    "role": "system",
                    "content": "ä½ æ˜¯ä¸€ä¸ªä¸“ä¸šçš„çŸ­è§†é¢‘å‰ªè¾‘ä¼˜åŒ–ä¸“å®¶ï¼Œæ“…é•¿ä¼˜åŒ–è„šæœ¬ä»¥æå‡å‰ªè¾‘æ•ˆæœã€‚"
                },
                {"role": "user", "content": prompt}
            ]
            
            _, response_text = ark_deepseek_v3_2(messages, max_new_tokens=4096)
            optimized_result = self._parse_json_response(response_text)
            
            if optimized_result:
                logger.info("Optimized script for editing")
                return optimized_result
            else:
                logger.warning("Failed to parse optimized script, returning original")
                return script_data
        except Exception as e:
            logger.error(f"Error optimizing script: {e}", exc_info=True)
            return script_data

